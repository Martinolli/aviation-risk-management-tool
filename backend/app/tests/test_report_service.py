import uuid
from datetime import datetime, timezone
from pathlib import Path
from zipfile import ZipFile

import pytest
from docx import Document
from sqlalchemy import create_engine, select
from sqlalchemy.orm import Session, sessionmaker
from sqlalchemy.pool import StaticPool

import app.models  # noqa: F401
from app.models.audit import AuditLog
from app.models.base import Base
from app.models.committee import Committee
from app.models.enums import (
    AuditAction,
    AuthorityLevel,
    CommitteeType,
    RiskActionStatus,
    RiskAssessmentType,
    RiskDecisionType,
    RiskDomain,
    RiskLifecycleStatus,
    RiskWorkflowStatus,
)
from app.models.risk import (
    RiskAction,
    RiskAssessment,
    RiskDecision,
    RiskEvidence,
    RiskRecord,
)
from app.services.report_service import (
    ReportRiskNotFoundError,
    generate_risk_dossier_docx,
)


@pytest.fixture()
def db_session() -> Session:
    engine = create_engine(
        "sqlite+pysqlite:///:memory:",
        connect_args={"check_same_thread": False},
        poolclass=StaticPool,
    )
    Base.metadata.create_all(engine)
    SessionLocal = sessionmaker(bind=engine)

    with SessionLocal() as session:
        yield session

    Base.metadata.drop_all(engine)


def _create_committee(db_session: Session) -> Committee:
    committee = Committee(
        name=f"Report Committee {uuid.uuid4()}",
        authority_level=AuthorityLevel.LOW,
        committee_type=CommitteeType.OPERATIONAL_BOARD,
        is_fixed=False,
        is_active=True,
    )
    db_session.add(committee)
    db_session.flush()
    return committee


def _create_risk_record(
    db_session: Session,
    *,
    risk_id: str | None = "RISK-2026-0001",
    board_of_origin_id: uuid.UUID | None = None,
) -> RiskRecord:
    risk_record = RiskRecord(
        risk_id=risk_id,
        problem_description="Unexpected vibration observed during taxi test.",
        source_trigger="Pilot report",
        domain=RiskDomain.FLIGHT_TEST,
        board_of_origin_id=board_of_origin_id,
        system_scope="Flight test aircraft",
        central_event="Vibration during ground movement",
        hazard_statement="Loss of component integrity could affect safety margin.",
        causes=["Loose instrumentation mount"],
        consequences=["Equipment damage"],
        existing_controls=["Pre-flight inspection"],
        workflow_status=RiskWorkflowStatus.APPROVED_AT_OPERATIONAL_BOARD,
        lifecycle_status=RiskLifecycleStatus.OPEN,
        is_active=True,
    )
    db_session.add(risk_record)
    db_session.flush()
    return risk_record


def _seed_report_detail(db_session: Session, risk_record: RiskRecord) -> None:
    now = datetime.now(timezone.utc)
    committee = _create_committee(db_session)
    db_session.add_all(
        [
            RiskAssessment(
                risk_record_id=risk_record.id,
                assessment_type=RiskAssessmentType.INITIAL,
                severity="Major",
                likelihood="Remote",
                risk_level="Medium",
                rationale="Initial safety risk assessment.",
                assessed_at=now,
            ),
            RiskAction(
                risk_record_id=risk_record.id,
                title="Inspect instrumentation mount",
                description="Verify mount installation and torque.",
                status=RiskActionStatus.COMPLETED,
                completion_notes="Mount inspected and secured.",
                completed_at=now,
            ),
            RiskDecision(
                risk_record_id=risk_record.id,
                committee_id=committee.id,
                decision_type=RiskDecisionType.APPROVE,
                decision_text="Operational board approved continued testing.",
                decided_at=now,
            ),
            AuditLog(
                entity_type="RiskRecord",
                entity_id=risk_record.id,
                action=AuditAction.CREATE,
                changed_at=now,
            ),
            AuditLog(
                entity_type="RiskRecord",
                entity_id=risk_record.id,
                action=AuditAction.UPDATE,
                changed_at=now,
            ),
            AuditLog(
                entity_type="RiskRecord",
                entity_id=risk_record.id,
                action=AuditAction.APPROVE,
                changed_at=now,
            ),
        ]
    )
    db_session.flush()


def _create_evidence(
    db_session: Session,
    risk_record: RiskRecord,
    *,
    filename: str = "flight-test-evidence.pdf",
    description: str = "Flight test supporting document",
    storage_path: str = "unused/report-metadata-only.pdf",
    file_size_bytes: int = 2048,
    is_active: bool = True,
) -> RiskEvidence:
    now = datetime.now(timezone.utc)
    evidence = RiskEvidence(
        risk_record_id=risk_record.id,
        original_filename=filename,
        stored_filename=f"{uuid.uuid4()}_{filename}",
        storage_path=storage_path,
        content_type="application/pdf",
        file_size_bytes=file_size_bytes,
        description=description,
        uploaded_at=now,
        is_active=is_active,
        archived_at=None if is_active else now,
        archive_reason=None if is_active else "Superseded evidence",
    )
    db_session.add(evidence)
    db_session.flush()
    db_session.add(
        AuditLog(
            entity_type="RiskEvidence",
            entity_id=evidence.id,
            action=AuditAction.CREATE if is_active else AuditAction.ARCHIVE,
            changed_at=now,
            new_value={"original_filename": filename},
            reason=None if is_active else "Superseded evidence",
        )
    )
    db_session.flush()
    return evidence


def _docx_text(path: Path) -> str:
    document = Document(path)
    pieces: list[str] = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                pieces.append(cell.text)
    return "\n".join(pieces)


def _generate_report(
    db_session: Session,
    tmp_path: Path,
    *,
    risk_id: str | None = "RISK-2026-0001",
) -> tuple[Path, RiskRecord]:
    risk_record = _create_risk_record(db_session, risk_id=risk_id)
    _seed_report_detail(db_session, risk_record)
    path = generate_risk_dossier_docx(
        db_session,
        risk_record_id=risk_record.id,
        output_dir=tmp_path / "reports",
    )
    return path, risk_record


def test_generate_risk_dossier_docx_creates_docx_file(
    db_session: Session,
    tmp_path: Path,
) -> None:
    path, _risk_record = _generate_report(db_session, tmp_path)

    assert path.exists()
    assert path.suffix == ".docx"


def test_generated_filename_uses_risk_id_when_available(
    db_session: Session,
    tmp_path: Path,
) -> None:
    path, _risk_record = _generate_report(db_session, tmp_path)

    assert path.name == "RISK-2026-0001_dossier.docx"


def test_generated_filename_falls_back_to_uuid_when_risk_id_missing(
    db_session: Session,
    tmp_path: Path,
) -> None:
    path, risk_record = _generate_report(db_session, tmp_path, risk_id=None)

    assert path.name == f"risk_{risk_record.id}_dossier.docx"


@pytest.mark.parametrize(
    "expected_text",
    [
        "Risk Dossier Report",
        "Unexpected vibration observed during taxi test.",
        "RISK-2026-0001",
        "Major",
        "Inspect instrumentation mount",
        "Operational board approved continued testing.",
        "Total Audit Records",
    ],
)
def test_generated_report_contains_expected_content(
    db_session: Session,
    tmp_path: Path,
    expected_text: str,
) -> None:
    path, _risk_record = _generate_report(db_session, tmp_path)

    assert expected_text in _docx_text(path)


def test_generated_report_contains_board_of_origin_traceability(
    db_session: Session,
    tmp_path: Path,
) -> None:
    board = _create_committee(db_session)
    risk_record = _create_risk_record(
        db_session,
        board_of_origin_id=board.id,
    )
    _seed_report_detail(db_session, risk_record)

    path = generate_risk_dossier_docx(
        db_session,
        risk_record_id=risk_record.id,
        output_dir=tmp_path / "reports",
    )
    report_text = _docx_text(path)

    assert "Board of Origin / Originating Committee" in report_text
    assert board.name in report_text
    assert f"Board of Origin ID: {board.id}" in report_text
    assert "Board of Origin Authority Level: Low" in report_text


def test_generated_report_contains_evidence_and_audit_annex_sections(
    db_session: Session,
    tmp_path: Path,
) -> None:
    path, _risk_record = _generate_report(db_session, tmp_path)
    report_text = _docx_text(path)

    assert "Section 6 - Evidence / Supporting Documents" in report_text
    assert "No evidence or supporting documents recorded." in report_text
    assert "Section 8 - Audit Trail Annex" in report_text
    assert "Create" in report_text
    assert "Section 9 - Notes / Disclaimer" in report_text


def test_generated_report_contains_active_and_archived_evidence_metadata(
    db_session: Session,
    tmp_path: Path,
) -> None:
    risk_record = _create_risk_record(db_session)
    _seed_report_detail(db_session, risk_record)
    active = _create_evidence(db_session, risk_record)
    archived = _create_evidence(
        db_session,
        risk_record,
        filename="archived-evidence.txt",
        description="Archived supporting note",
        is_active=False,
    )

    path = generate_risk_dossier_docx(
        db_session,
        risk_record_id=risk_record.id,
        output_dir=tmp_path / "reports",
    )
    report_text = _docx_text(path)

    assert active.original_filename in report_text
    assert active.description in report_text
    assert archived.original_filename in report_text
    assert archived.description in report_text
    assert "2.0 KB" in report_text
    assert "Active" in report_text
    assert "Archived" in report_text
    assert "Superseded evidence" in report_text
    assert "Evidence Records: 2" in report_text
    assert "RiskEvidence" in report_text
    assert "Archive" in report_text


def test_generated_report_uses_committee_name_and_authority_level(
    db_session: Session,
    tmp_path: Path,
) -> None:
    path, _risk_record = _generate_report(db_session, tmp_path)
    committee = db_session.scalar(select(Committee))
    assert committee is not None
    report_text = _docx_text(path)

    assert committee.name in report_text
    assert "Authority Level" in report_text
    assert "Low" in report_text


def test_generated_report_with_only_create_audit_still_generates(
    db_session: Session,
    tmp_path: Path,
) -> None:
    risk_record = _create_risk_record(db_session)
    db_session.add(
        AuditLog(
            entity_type="RiskRecord",
            entity_id=risk_record.id,
            action=AuditAction.CREATE,
            changed_at=datetime.now(timezone.utc),
        )
    )
    db_session.flush()

    path = generate_risk_dossier_docx(
        db_session,
        risk_record_id=risk_record.id,
        output_dir=tmp_path / "reports",
    )

    assert path.is_file()
    assert "Total Audit Records: 1" in _docx_text(path)


def test_evidence_file_binary_is_not_embedded_in_docx(
    db_session: Session,
    tmp_path: Path,
) -> None:
    risk_record = _create_risk_record(db_session)
    payload = b"TASK074_UNIQUE_EVIDENCE_BINARY_PAYLOAD"
    evidence_path = tmp_path / "source-evidence.bin"
    evidence_path.write_bytes(payload)
    _create_evidence(
        db_session,
        risk_record,
        filename=evidence_path.name,
        storage_path=str(evidence_path),
        file_size_bytes=len(payload),
    )

    path = generate_risk_dossier_docx(
        db_session,
        risk_record_id=risk_record.id,
        output_dir=tmp_path / "reports",
    )

    with ZipFile(path) as archive:
        member_names = archive.namelist()
        assert not any(name.startswith("word/embeddings/") for name in member_names)
        assert not any(payload in archive.read(name) for name in member_names)


def test_report_tables_use_grid_style_and_bold_headers(
    db_session: Session,
    tmp_path: Path,
) -> None:
    path, _risk_record = _generate_report(db_session, tmp_path)
    document = Document(path)
    assessment_table = next(
        table for table in document.tables if table.cell(0, 0).text == "Type"
    )

    assert assessment_table.style.name == "Table Grid"
    assert all(
        run.bold
        for cell in assessment_table.rows[0].cells
        for run in cell.paragraphs[0].runs
    )


def test_generating_report_for_unknown_risk_raises_not_found(
    db_session: Session,
    tmp_path: Path,
) -> None:
    with pytest.raises(ReportRiskNotFoundError):
        generate_risk_dossier_docx(
            db_session,
            risk_record_id=uuid.uuid4(),
            output_dir=tmp_path / "reports",
        )


def test_output_dir_is_created_if_it_does_not_exist(
    db_session: Session,
    tmp_path: Path,
) -> None:
    output_dir = tmp_path / "nested" / "reports"
    risk_record = _create_risk_record(db_session)
    _seed_report_detail(db_session, risk_record)

    path = generate_risk_dossier_docx(
        db_session,
        risk_record_id=risk_record.id,
        output_dir=output_dir,
    )

    assert output_dir.exists()
    assert path.exists()

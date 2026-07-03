import uuid
from collections.abc import Generator
from datetime import date, datetime, timedelta, timezone
from pathlib import Path

import pytest
from docx import Document
from fastapi.testclient import TestClient
from sqlalchemy import create_engine, select
from sqlalchemy.orm import Session, sessionmaker
from sqlalchemy.pool import StaticPool

import app.models  # noqa: F401
from app.core.database import get_db
from app.main import app
from app.models.audit import AuditLog
from app.models.base import Base
from app.models.committee import Committee, CommitteeMember
from app.models.enums import (
    AuditAction,
    AuthorityLevel,
    CommitteeType,
    RiskActionStatus,
    RiskAssessmentType,
    RiskDomain,
    RiskLifecycleStatus,
    RiskMonitoringStatus,
    RiskWorkflowStatus,
)
from app.models.risk import (
    RiskAction,
    RiskAssessment,
    RiskEvidence,
    RiskMonitoringReview,
    RiskRecord,
)
from app.models.user import User
from app.services.committee_meeting_pack_service import (
    CommitteeMeetingPackBusinessRuleError,
    generate_committee_meeting_pack_docx,
)
from app.services.report_tracking_service import (
    COMMITTEE_MEETING_PACK_REPORT_TYPE,
    generate_and_track_committee_meeting_pack,
)


@pytest.fixture()
def db_session() -> Generator[Session, None, None]:
    engine = create_engine(
        "sqlite+pysqlite:///:memory:",
        connect_args={"check_same_thread": False},
        poolclass=StaticPool,
    )
    Base.metadata.create_all(engine)
    SessionLocal = sessionmaker(bind=engine, autoflush=False, autocommit=False)
    with SessionLocal() as session:
        yield session
    Base.metadata.drop_all(engine)


@pytest.fixture()
def client(db_session: Session) -> Generator[TestClient, None, None]:
    def override_get_db() -> Generator[Session, None, None]:
        yield db_session

    app.dependency_overrides[get_db] = override_get_db
    with TestClient(app) as test_client:
        yield test_client
    app.dependency_overrides.clear()


def _create_user(db: Session) -> User:
    user = User(
        email=f"meeting-pack-{uuid.uuid4()}@example.com",
        display_name="Meeting Pack User",
        is_active=True,
    )
    db.add(user)
    db.flush()
    return user


def _create_committee(
    db: Session,
    *,
    authority_level: AuthorityLevel = AuthorityLevel.LOW,
    name: str | None = None,
) -> Committee:
    committee_type = {
        AuthorityLevel.LOW: CommitteeType.OPERATIONAL_BOARD,
        AuthorityLevel.MIDDLE: CommitteeType.RISK_MANAGEMENT_COMMITTEE,
        AuthorityLevel.HIGH: CommitteeType.EXECUTIVE_SAFETY_MANAGEMENT_COMMITTEE,
    }[authority_level]
    committee = Committee(
        name=name or f"{authority_level.value} Meeting Committee {uuid.uuid4()}",
        authority_level=authority_level,
        committee_type=committee_type,
        is_fixed=authority_level in {AuthorityLevel.MIDDLE, AuthorityLevel.HIGH},
        is_active=True,
    )
    db.add(committee)
    db.flush()
    return committee


def _add_member(db: Session, committee: Committee, user: User) -> CommitteeMember:
    member = CommitteeMember(
        committee_id=committee.id,
        user_id=user.id,
        role_label="Committee Member",
        is_active=True,
    )
    db.add(member)
    db.flush()
    return member


def _queue_status(authority_level: AuthorityLevel) -> RiskWorkflowStatus:
    return {
        AuthorityLevel.LOW: RiskWorkflowStatus.SUBMITTED_TO_OPERATIONAL_BOARD,
        AuthorityLevel.MIDDLE: RiskWorkflowStatus.ESCALATED_TO_RISK_MANAGEMENT_COMMITTEE,
        AuthorityLevel.HIGH: RiskWorkflowStatus.ESCALATED_TO_EXECUTIVE_COMMITTEE,
    }[authority_level]


def _create_risk(
    db: Session,
    *,
    committee: Committee,
    workflow_status: RiskWorkflowStatus | None = None,
    risk_id: str | None = None,
) -> RiskRecord:
    risk = RiskRecord(
        risk_id=risk_id or f"RISK-{uuid.uuid4().hex[:8]}",
        problem_description="Queued flight test control risk",
        domain=RiskDomain.FLIGHT_TEST,
        board_of_origin_id=committee.id,
        system_scope="Flight test instrumentation",
        central_event="Loss of instrumentation data",
        hazard_statement="Instrumentation failure may affect test safety",
        workflow_status=workflow_status or _queue_status(committee.authority_level),
        lifecycle_status=RiskLifecycleStatus.OPEN,
        is_active=True,
    )
    db.add(risk)
    db.flush()
    return risk


def _seed_risk_context(db: Session, risk: RiskRecord) -> None:
    db.add_all(
        [
            RiskAssessment(
                risk_record_id=risk.id,
                assessment_type=RiskAssessmentType.INITIAL,
                severity="Major",
                likelihood="Possible",
                risk_level="High",
                assessed_at=datetime.now(timezone.utc),
            ),
            RiskAction(
                risk_record_id=risk.id,
                title="Inspect instrumentation wiring",
                action_owner_user_id=uuid.uuid4(),
                due_date=date.today() - timedelta(days=2),
                status=RiskActionStatus.OPEN,
            ),
            RiskEvidence(
                risk_record_id=risk.id,
                original_filename="inspection-plan.pdf",
                stored_filename="inspection-plan.pdf",
                storage_path="evidence/inspection-plan.pdf",
                file_size_bytes=1024,
                description="Inspection plan metadata",
                uploaded_at=datetime.now(timezone.utc),
                is_active=True,
            ),
            RiskMonitoringReview(
                risk_record_id=risk.id,
                review_frequency="Monthly",
                next_review_date=date.today(),
                status=RiskMonitoringStatus.DUE,
                is_active=True,
            ),
        ]
    )
    db.flush()


def _docx_text(path: Path) -> str:
    document = Document(path)
    text = [paragraph.text for paragraph in document.paragraphs]
    text.extend(cell.text for table in document.tables for row in table.rows for cell in row.cells)
    return "\n".join(text)


def _headers(user: User) -> dict[str, str]:
    return {"X-User-Id": str(user.id)}


def test_authorized_member_generates_tracked_pack_with_required_content(
    db_session: Session,
    tmp_path: Path,
) -> None:
    member = _create_user(db_session)
    committee = _create_committee(
        db_session,
        name="Flight Test Safety Committee - Operation",
    )
    _add_member(db_session, committee, member)
    risk = _create_risk(db_session, committee=committee, risk_id="RISK-PACK-001")
    _seed_risk_context(db_session, risk)

    report = generate_and_track_committee_meeting_pack(
        db_session,
        committee_id=committee.id,
        output_dir=tmp_path,
        generated_by_user_id=member.id,
        meeting_title="Monthly Safety Review",
        meeting_date=date(2026, 7, 15),
    )
    text = _docx_text(Path(report.file_path))
    audit_log = db_session.scalar(
        select(AuditLog).where(
            AuditLog.entity_type == "Committee",
            AuditLog.entity_id == committee.id,
            AuditLog.action == AuditAction.GENERATE_REPORT,
        )
    )

    assert report.report_type == COMMITTEE_MEETING_PACK_REPORT_TYPE
    assert report.committee_id == committee.id
    assert report.risk_record_id is None
    assert Path(report.file_path).is_file()
    for expected in (
        "Committee Meeting Pack",
        committee.name,
        "Authority Level",
        "Meeting Agenda",
        "Decision Queue Summary",
        "Risks Awaiting Committee Review",
        "RISK-PACK-001",
        risk.problem_description,
        "Action Follow-up",
        "Overdue",
        "Monitoring Follow-up",
        "Evidence Summary",
        "inspection-plan.pdf",
    ):
        assert expected in text
    assert audit_log is not None
    assert audit_log.changed_by_user_id == member.id


def test_empty_decision_queue_still_generates_agenda(
    db_session: Session,
    tmp_path: Path,
) -> None:
    member = _create_user(db_session)
    committee = _create_committee(db_session)
    _add_member(db_session, committee, member)

    path = generate_committee_meeting_pack_docx(
        db_session,
        committee_id=committee.id,
        generated_by_user_id=member.id,
        output_dir=tmp_path,
    )
    text = _docx_text(path)

    assert "Meeting Agenda" in text
    assert "Total risks awaiting committee decision: 0" in text
    assert "No risks are currently in this committee Decision Queue." in text


def test_unauthorized_user_cannot_generate_pack(
    db_session: Session,
    tmp_path: Path,
) -> None:
    member = _create_user(db_session)
    unrelated = _create_user(db_session)
    committee = _create_committee(db_session)
    _add_member(db_session, committee, member)

    with pytest.raises(CommitteeMeetingPackBusinessRuleError, match="not authorized"):
        generate_committee_meeting_pack_docx(
            db_session,
            committee_id=committee.id,
            generated_by_user_id=unrelated.id,
            output_dir=tmp_path,
        )


@pytest.mark.parametrize(
    "authority_level",
    [AuthorityLevel.LOW, AuthorityLevel.MIDDLE, AuthorityLevel.HIGH],
)
def test_pack_uses_authority_specific_decision_queue(
    db_session: Session,
    tmp_path: Path,
    authority_level: AuthorityLevel,
) -> None:
    member = _create_user(db_session)
    committee = _create_committee(db_session, authority_level=authority_level)
    other_committee = _create_committee(db_session)
    _add_member(db_session, committee, member)
    included = _create_risk(
        db_session,
        committee=committee,
        risk_id=f"RISK-{authority_level.value}-INCLUDED",
    )
    excluded = _create_risk(
        db_session,
        committee=other_committee,
        workflow_status=RiskWorkflowStatus.SUBMITTED_TO_OPERATIONAL_BOARD,
        risk_id=f"RISK-{authority_level.value}-EXCLUDED",
    )

    path = generate_committee_meeting_pack_docx(
        db_session,
        committee_id=committee.id,
        generated_by_user_id=member.id,
        output_dir=tmp_path / authority_level.value,
    )
    text = _docx_text(path)

    assert included.risk_id in text
    assert excluded.risk_id not in text


def test_committee_report_api_supports_generate_list_get_and_download(
    client: TestClient,
    db_session: Session,
    tmp_path: Path,
) -> None:
    member = _create_user(db_session)
    unrelated = _create_user(db_session)
    committee = _create_committee(db_session)
    _add_member(db_session, committee, member)
    _create_risk(db_session, committee=committee)

    created = client.post(
        f"/reports/committee-meeting-packs/{committee.id}",
        headers=_headers(member),
        json={
            "output_dir": str(tmp_path),
            "meeting_title": "Committee Review",
            "meeting_date": "2026-07-20",
        },
    )

    assert created.status_code == 201
    report = created.json()
    assert report["report_type"] == COMMITTEE_MEETING_PACK_REPORT_TYPE
    assert report["committee_id"] == str(committee.id)
    assert report["risk_record_id"] is None

    authorized_list = client.get("/reports", headers=_headers(member))
    authorized_get = client.get(
        f"/reports/{report['id']}", headers=_headers(member)
    )
    authorized_download = client.get(
        f"/reports/{report['id']}/download", headers=_headers(member)
    )
    unauthorized_list = client.get("/reports", headers=_headers(unrelated))
    unauthorized_download = client.get(
        f"/reports/{report['id']}/download", headers=_headers(unrelated)
    )

    assert [item["id"] for item in authorized_list.json()] == [report["id"]]
    assert authorized_get.status_code == 200
    assert authorized_download.status_code == 200
    assert authorized_download.content
    assert unauthorized_list.json() == []
    assert unauthorized_download.status_code == 400

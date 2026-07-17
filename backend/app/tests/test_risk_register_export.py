import csv
import io
import uuid
from collections.abc import Generator
from datetime import date, datetime, timedelta, timezone
from pathlib import Path

import pytest
from docx import Document
from fastapi.testclient import TestClient
from sqlalchemy import create_engine
from sqlalchemy.orm import Session, sessionmaker
from sqlalchemy.pool import StaticPool

import app.models  # noqa: F401
from app.core.database import get_db
from app.main import app
from app.models.base import Base
from app.models.committee import Committee, CommitteeMember
from app.models.enums import (
    AuthorityLevel,
    CommitteeType,
    RiskActionStatus,
    RiskAssessmentType,
    RiskDomain,
    RiskLifecycleStatus,
    RiskMonitoringStatus,
    RiskWorkflowStatus,
)
from app.models.risk import (
    RiskAction,
    RiskAssessment,
    RiskEvidence,
    RiskMonitoringReview,
    RiskRecord,
)
from app.models.user import User
from app.services.auth_service import create_access_token
from app.services.risk_register_export_service import RISK_REGISTER_EXPORT_COLUMNS


@pytest.fixture()
def db_session() -> Generator[Session, None, None]:
    engine = create_engine(
        "sqlite+pysqlite:///:memory:",
        connect_args={"check_same_thread": False},
        poolclass=StaticPool,
    )
    Base.metadata.create_all(engine)
    SessionLocal = sessionmaker(bind=engine, autoflush=False, autocommit=False)

    with SessionLocal() as session:
        yield session

    Base.metadata.drop_all(engine)


@pytest.fixture()
def client(
    db_session: Session,
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> Generator[TestClient, None, None]:
    monkeypatch.setattr("app.api.risks.RISK_REGISTER_EXPORT_OUTPUT_DIR", tmp_path)

    def override_get_db() -> Generator[Session, None, None]:
        yield db_session

    app.dependency_overrides[get_db] = override_get_db
    with TestClient(app) as test_client:
        yield test_client
    app.dependency_overrides.clear()


def _create_user(db: Session, *, is_active: bool = True) -> User:
    user = User(
        email=f"risk-register-export-{uuid.uuid4()}@example.com",
        display_name="Risk Register Export User",
        is_active=is_active,
    )
    db.add(user)
    db.commit()
    db.refresh(user)
    return user


def _headers(user: User) -> dict[str, str]:
    return {"Authorization": f"Bearer {create_access_token(user_id=user.id)}"}


def _create_board(db: Session, *, name: str = "Export Board") -> Committee:
    board = Committee(
        name=f"{name} {uuid.uuid4()}",
        authority_level=AuthorityLevel.LOW,
        committee_type=CommitteeType.OPERATIONAL_BOARD,
        is_fixed=False,
        is_active=True,
    )
    db.add(board)
    db.commit()
    db.refresh(board)
    return board


def _create_fixed_governance_committee(
    db: Session,
    *,
    authority_level: AuthorityLevel,
) -> Committee:
    committee = Committee(
        name=f"{authority_level.value} Export Governance {uuid.uuid4()}",
        authority_level=authority_level,
        committee_type=(
            CommitteeType.RISK_MANAGEMENT_COMMITTEE
            if authority_level == AuthorityLevel.MIDDLE
            else CommitteeType.EXECUTIVE_SAFETY_MANAGEMENT_COMMITTEE
        ),
        is_fixed=True,
        is_active=True,
    )
    db.add(committee)
    db.commit()
    db.refresh(committee)
    return committee


def _add_membership(db: Session, *, committee: Committee, user: User) -> None:
    db.add(
        CommitteeMember(
            committee_id=committee.id,
            user_id=user.id,
            is_active=True,
        )
    )
    db.commit()


def _create_risk(
    db: Session,
    *,
    creator: User,
    risk_id: str | None = None,
    problem_description: str = "Risk register export hydraulic event",
    source_trigger: str | None = "Voluntary safety report",
    domain: RiskDomain = RiskDomain.FLIGHT_TEST,
    board_of_origin_id: uuid.UUID | None = None,
    workflow_status: RiskWorkflowStatus = RiskWorkflowStatus.DRAFT,
    lifecycle_status: RiskLifecycleStatus = RiskLifecycleStatus.OPEN,
    is_active: bool = True,
) -> RiskRecord:
    risk = RiskRecord(
        risk_id=risk_id or f"RISK-EXPORT-{uuid.uuid4()}",
        problem_description=problem_description,
        source_trigger=source_trigger,
        domain=domain,
        board_of_origin_id=board_of_origin_id,
        workflow_status=workflow_status,
        lifecycle_status=lifecycle_status,
        created_by_user_id=creator.id,
        is_active=is_active,
        archived_at=None if is_active else datetime.now(timezone.utc),
    )
    db.add(risk)
    db.commit()
    db.refresh(risk)
    return risk


def _add_assessment(
    db: Session,
    risk: RiskRecord,
    *,
    risk_level: str,
    assessment_type: RiskAssessmentType = RiskAssessmentType.INITIAL,
    assessed_at: datetime | None = None,
) -> RiskAssessment:
    assessment = RiskAssessment(
        risk_record_id=risk.id,
        assessment_type=assessment_type,
        severity="Major",
        likelihood="Remote",
        risk_level=risk_level,
        assessed_at=assessed_at or datetime.now(timezone.utc),
    )
    db.add(assessment)
    db.commit()
    db.refresh(assessment)
    return assessment


def _add_action(
    db: Session,
    risk: RiskRecord,
    *,
    status: RiskActionStatus = RiskActionStatus.OPEN,
    due_date: date | None = None,
) -> None:
    db.add(
        RiskAction(
            risk_record_id=risk.id,
            title=f"Action {uuid.uuid4()}",
            due_date=due_date,
            status=status,
        )
    )
    db.commit()


def _add_monitoring(
    db: Session,
    risk: RiskRecord,
    *,
    status: RiskMonitoringStatus,
    is_active: bool = True,
) -> None:
    db.add(
        RiskMonitoringReview(
            risk_record_id=risk.id,
            status=status,
            is_active=is_active,
        )
    )
    db.commit()


def _add_evidence(db: Session, risk: RiskRecord, *, is_active: bool = True) -> None:
    now = datetime.now(timezone.utc)
    db.add(
        RiskEvidence(
            risk_record_id=risk.id,
            original_filename="risk-evidence.pdf",
            stored_filename=f"{uuid.uuid4()}_risk-evidence.pdf",
            storage_path="unused/risk-evidence.pdf",
            file_size_bytes=1024,
            uploaded_at=now,
            is_active=is_active,
            archived_at=None if is_active else now,
        )
    )
    db.commit()


def _csv_rows(response) -> list[dict[str, str]]:
    content = response.content.decode("utf-8")
    return list(csv.DictReader(io.StringIO(content)))


def _docx_text(response) -> str:
    document = Document(io.BytesIO(response.content))
    parts: list[str] = []
    parts.extend(paragraph.text for paragraph in document.paragraphs)
    for table in document.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


def test_authorized_user_can_export_csv_and_media_type(
    client: TestClient, db_session: Session
) -> None:
    user = _create_user(db_session)
    _create_risk(db_session, creator=user)

    response = client.get("/risks/export/csv", headers=_headers(user))

    assert response.status_code == 200
    assert response.headers["content-type"].startswith("text/csv")


def test_csv_contains_expected_headers(
    client: TestClient, db_session: Session
) -> None:
    user = _create_user(db_session)
    _create_risk(db_session, creator=user)

    response = client.get("/risks/export/csv", headers=_headers(user))
    header = response.content.decode("utf-8").splitlines()[0]

    assert header.split(",") == RISK_REGISTER_EXPORT_COLUMNS


def test_csv_contains_authorized_records_and_excludes_unauthorized(
    client: TestClient, db_session: Session
) -> None:
    user = _create_user(db_session)
    other = _create_user(db_session)
    readable = _create_risk(db_session, creator=user, risk_id="RISK-AUTHORIZED")
    hidden = _create_risk(db_session, creator=other, risk_id="RISK-HIDDEN")

    response = client.get("/risks/export/csv", headers=_headers(user))
    exported_ids = {row["risk_record_id"] for row in _csv_rows(response)}

    assert str(readable.id) in exported_ids
    assert str(hidden.id) not in exported_ids


def test_fixed_middle_high_governance_users_export_oversight_records(
    client: TestClient,
    db_session: Session,
) -> None:
    middle_user = _create_user(db_session)
    high_user = _create_user(db_session)
    risk_owner = _create_user(db_session)
    middle_committee = _create_fixed_governance_committee(
        db_session,
        authority_level=AuthorityLevel.MIDDLE,
    )
    high_committee = _create_fixed_governance_committee(
        db_session,
        authority_level=AuthorityLevel.HIGH,
    )
    _add_membership(db_session, committee=middle_committee, user=middle_user)
    _add_membership(db_session, committee=high_committee, user=high_user)
    oversight_risk = _create_risk(
        db_session,
        creator=risk_owner,
        risk_id="RISK-GOVERNANCE-OVERSIGHT",
    )

    middle_response = client.get("/risks/export/csv", headers=_headers(middle_user))
    high_response = client.get("/risks/export/csv", headers=_headers(high_user))

    assert str(oversight_risk.id) in {
        row["risk_record_id"] for row in _csv_rows(middle_response)
    }
    assert str(oversight_risk.id) in {
        row["risk_record_id"] for row in _csv_rows(high_response)
    }


@pytest.mark.parametrize(
    ("query", "target_kwargs", "other_kwargs"),
    [
        (
            "search=hydraulic",
            {"problem_description": "Hydraulic pressure drop"},
            {"problem_description": "Cabin placard issue"},
        ),
        (
            "risk_id=0042",
            {"risk_id": "RISK-2026-0042"},
            {"risk_id": "RISK-2026-0099"},
        ),
        (
            "domain=ENGINEERING",
            {"domain": RiskDomain.ENGINEERING},
            {"domain": RiskDomain.QUALITY},
        ),
        (
            "workflow_status=ACCEPTED",
            {"workflow_status": RiskWorkflowStatus.ACCEPTED},
            {"workflow_status": RiskWorkflowStatus.DRAFT},
        ),
        (
            "lifecycle_status=MONITORING",
            {"lifecycle_status": RiskLifecycleStatus.MONITORING},
            {"lifecycle_status": RiskLifecycleStatus.OPEN},
        ),
    ],
)
def test_csv_respects_basic_filters(
    client: TestClient,
    db_session: Session,
    query: str,
    target_kwargs: dict[str, object],
    other_kwargs: dict[str, object],
) -> None:
    user = _create_user(db_session)
    target = _create_risk(db_session, creator=user, **target_kwargs)
    _create_risk(db_session, creator=user, **other_kwargs)

    response = client.get(f"/risks/export/csv?{query}", headers=_headers(user))

    assert [row["risk_record_id"] for row in _csv_rows(response)] == [str(target.id)]


def test_csv_respects_board_of_origin_filter(
    client: TestClient, db_session: Session
) -> None:
    user = _create_user(db_session)
    board = _create_board(db_session, name="Target Board")
    other_board = _create_board(db_session, name="Other Board")
    target = _create_risk(db_session, creator=user, board_of_origin_id=board.id)
    _create_risk(db_session, creator=user, board_of_origin_id=other_board.id)

    response = client.get(
        f"/risks/export/csv?board_of_origin_id={board.id}",
        headers=_headers(user),
    )
    row = _csv_rows(response)[0]

    assert row["risk_record_id"] == str(target.id)
    assert row["board_of_origin_name"] == board.name
    assert row["board_of_origin_authority_level"] == "LOW"


def test_csv_respects_latest_risk_level_filter_and_includes_latest_assessment(
    client: TestClient, db_session: Session
) -> None:
    user = _create_user(db_session)
    older_high_latest_low = _create_risk(db_session, creator=user)
    latest_high = _create_risk(db_session, creator=user)
    now = datetime.now(timezone.utc)
    _add_assessment(db_session, older_high_latest_low, risk_level="HIGH", assessed_at=now)
    _add_assessment(
        db_session,
        older_high_latest_low,
        risk_level="LOW",
        assessment_type=RiskAssessmentType.RESIDUAL,
        assessed_at=now + timedelta(hours=1),
    )
    _add_assessment(
        db_session,
        latest_high,
        risk_level="HIGH",
        assessment_type=RiskAssessmentType.RESIDUAL,
        assessed_at=now,
    )

    response = client.get(
        "/risks/export/csv?latest_risk_level=HIGH",
        headers=_headers(user),
    )
    rows = _csv_rows(response)

    assert [row["risk_record_id"] for row in rows] == [str(latest_high.id)]
    assert rows[0]["latest_assessment_type"] == "RESIDUAL"
    assert rows[0]["latest_risk_level"] == "HIGH"
    assert rows[0]["latest_assessment_date"]


def test_csv_respects_overdue_action_filter_and_counts_actions(
    client: TestClient, db_session: Session
) -> None:
    user = _create_user(db_session)
    overdue = _create_risk(db_session, creator=user)
    current = _create_risk(db_session, creator=user)
    _add_action(
        db_session,
        overdue,
        status=RiskActionStatus.OPEN,
        due_date=date.today() - timedelta(days=1),
    )
    _add_action(
        db_session,
        overdue,
        status=RiskActionStatus.IN_PROGRESS,
        due_date=date.today() + timedelta(days=1),
    )
    _add_action(
        db_session,
        overdue,
        status=RiskActionStatus.COMPLETED,
        due_date=date.today() - timedelta(days=2),
    )
    _add_action(
        db_session,
        current,
        status=RiskActionStatus.OPEN,
        due_date=date.today() + timedelta(days=1),
    )

    response = client.get(
        "/risks/export/csv?has_overdue_actions=true",
        headers=_headers(user),
    )
    row = _csv_rows(response)[0]

    assert row["risk_record_id"] == str(overdue.id)
    assert row["open_action_count"] == "2"
    assert row["overdue_action_count"] == "1"


def test_csv_respects_monitoring_filter_and_status_priority(
    client: TestClient, db_session: Session
) -> None:
    user = _create_user(db_session)
    monitored = _create_risk(db_session, creator=user)
    not_due = _create_risk(db_session, creator=user)
    _add_monitoring(db_session, monitored, status=RiskMonitoringStatus.ACTIVE)
    _add_monitoring(db_session, monitored, status=RiskMonitoringStatus.DUE)
    _add_monitoring(db_session, monitored, status=RiskMonitoringStatus.OVERDUE)
    _add_monitoring(db_session, not_due, status=RiskMonitoringStatus.ACTIVE)

    response = client.get(
        "/risks/export/csv?has_due_or_overdue_monitoring=true",
        headers=_headers(user),
    )
    row = _csv_rows(response)[0]

    assert row["risk_record_id"] == str(monitored.id)
    assert row["monitoring_status"] == "OVERDUE"


def test_csv_respects_include_archived_flags(
    client: TestClient, db_session: Session
) -> None:
    user = _create_user(db_session)
    active = _create_risk(db_session, creator=user, is_active=True)
    archived = _create_risk(db_session, creator=user, is_active=False)

    default_response = client.get("/risks/export/csv", headers=_headers(user))
    archived_response = client.get(
        "/risks/export/csv?include_archived=true",
        headers=_headers(user),
    )

    default_ids = {row["risk_record_id"] for row in _csv_rows(default_response)}
    archived_ids = {row["risk_record_id"] for row in _csv_rows(archived_response)}
    assert default_ids == {str(active.id)}
    assert archived_ids == {str(active.id), str(archived.id)}


def test_csv_includes_monitoring_evidence_and_not_assessed_defaults(
    client: TestClient, db_session: Session
) -> None:
    user = _create_user(db_session)
    risk = _create_risk(db_session, creator=user)
    _add_monitoring(db_session, risk, status=RiskMonitoringStatus.CLOSED)
    _add_evidence(db_session, risk, is_active=True)
    _add_evidence(db_session, risk, is_active=False)

    response = client.get("/risks/export/csv", headers=_headers(user))
    row = _csv_rows(response)[0]

    assert row["latest_assessment_type"] == ""
    assert row["latest_risk_level"] == "Not assessed"
    assert row["monitoring_status"] == "CLOSED"
    assert row["evidence_count"] == "1"


def test_authorized_user_can_export_docx_and_media_type(
    client: TestClient, db_session: Session
) -> None:
    user = _create_user(db_session)
    _create_risk(db_session, creator=user)

    response = client.get("/risks/export/docx", headers=_headers(user))

    assert response.status_code == 200
    assert response.headers["content-type"] == (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )


def test_docx_contains_controlled_export_content_and_excludes_unauthorized(
    client: TestClient, db_session: Session
) -> None:
    user = _create_user(db_session)
    other = _create_user(db_session)
    readable = _create_risk(db_session, creator=user, risk_id="RISK-DOCX-READABLE")
    hidden = _create_risk(db_session, creator=other, risk_id="RISK-DOCX-HIDDEN")

    response = client.get("/risks/export/docx", headers=_headers(user))
    text = _docx_text(response)

    for expected in [
        "Risk Register Export",
        "SMS Risk Management Process Tool",
        "Controlled Export",
        "Register Summary",
        "Risk Register",
        "Risk Details",
        "Disclaimer",
        "SMS governance",
        readable.risk_id,
    ]:
        assert expected in text
    assert hidden.risk_id not in text


def test_invalid_sort_parameters_return_http_400(
    client: TestClient, db_session: Session
) -> None:
    user = _create_user(db_session)

    bad_sort_by = client.get(
        "/risks/export/csv?sort_by=owner_user_id",
        headers=_headers(user),
    )
    bad_direction = client.get(
        "/risks/export/csv?sort_direction=sideways",
        headers=_headers(user),
    )

    assert bad_sort_by.status_code == 400
    assert "Unsupported sort_by" in bad_sort_by.json()["error"]["message"]
    assert bad_direction.status_code == 400
    assert "Unsupported sort_direction" in bad_direction.json()["error"]["message"]


def test_unauthenticated_and_inactive_users_cannot_export(
    client: TestClient, db_session: Session
) -> None:
    inactive = _create_user(db_session, is_active=False)

    unauthenticated_response = client.get("/risks/export/csv")
    inactive_response = client.get("/risks/export/csv", headers=_headers(inactive))

    assert unauthenticated_response.status_code == 400
    assert inactive_response.status_code == 403


def test_existing_risk_list_and_detail_routes_still_work(
    client: TestClient, db_session: Session
) -> None:
    user = _create_user(db_session)
    risk = _create_risk(db_session, creator=user)

    list_response = client.get("/risks", headers=_headers(user))
    detail_response = client.get(f"/risks/{risk.id}", headers=_headers(user))

    assert list_response.status_code == 200
    assert detail_response.status_code == 200
    assert detail_response.json()["id"] == str(risk.id)


def test_export_files_are_created_in_tmp_path(
    client: TestClient, db_session: Session, tmp_path: Path
) -> None:
    user = _create_user(db_session)
    _create_risk(db_session, creator=user)

    csv_response = client.get("/risks/export/csv", headers=_headers(user))
    docx_response = client.get("/risks/export/docx", headers=_headers(user))

    assert csv_response.status_code == 200
    assert docx_response.status_code == 200
    assert list(tmp_path.glob("risk_register_export_*.csv"))
    assert list(tmp_path.glob("risk_register_export_*.docx"))

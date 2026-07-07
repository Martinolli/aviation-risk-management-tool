import uuid
from collections.abc import Generator
from datetime import datetime, timezone
from pathlib import Path

import pytest
from docx import Document
from fastapi.testclient import TestClient
from sqlalchemy import create_engine
from sqlalchemy.orm import Session, sessionmaker
from sqlalchemy.pool import StaticPool

import app.models  # noqa: F401
from app.core.database import get_db
from app.main import app
from app.models.base import Base
from app.models.committee import Committee, CommitteeMember
from app.models.enums import (
    AuthorityLevel,
    CommitteeType,
    RiskDecisionType,
    RiskDomain,
    RiskLifecycleStatus,
    RiskWorkflowStatus,
)
from app.models.risk import RiskDecision, RiskRecord
from app.models.user import User
from app.schemas.committee_meeting import CommitteeMeetingCreate
from app.services.committee_meeting_service import finalize_committee_meeting
from app.schemas.committee_meeting import CommitteeMeetingFinalize
from app.services.report_tracking_service import (
    COMMITTEE_MEETING_MINUTES_REPORT_TYPE,
    ReportTrackingBusinessRuleError,
    generate_and_track_committee_meeting_minutes_report,
)


@pytest.fixture()
def db_session() -> Generator[Session, None, None]:
    engine = create_engine(
        "sqlite+pysqlite:///:memory:",
        connect_args={"check_same_thread": False},
        poolclass=StaticPool,
    )
    Base.metadata.create_all(engine)
    SessionLocal = sessionmaker(bind=engine, autoflush=False, autocommit=False)
    with SessionLocal() as session:
        yield session
    Base.metadata.drop_all(engine)


@pytest.fixture()
def client(db_session: Session) -> Generator[TestClient, None, None]:
    def override_get_db() -> Generator[Session, None, None]:
        yield db_session

    app.dependency_overrides[get_db] = override_get_db
    with TestClient(app) as test_client:
        yield test_client
    app.dependency_overrides.clear()


def _headers(user: User) -> dict[str, str]:
    return {"X-User-Id": str(user.id)}


def _docx_text(path: Path) -> str:
    document = Document(path)
    text = [paragraph.text for paragraph in document.paragraphs]
    text.extend(cell.text for table in document.tables for row in table.rows for cell in row.cells)
    return "\n".join(text)


def _create_user(db: Session) -> User:
    user = User(
        email=f"minutes-report-{uuid.uuid4()}@example.com",
        display_name="Minutes Report User",
        is_active=True,
    )
    db.add(user)
    db.flush()
    return user


def _create_committee(db: Session) -> Committee:
    committee = Committee(
        name=f"Minutes Report Committee {uuid.uuid4()}",
        authority_level=AuthorityLevel.LOW,
        committee_type=CommitteeType.OPERATIONAL_BOARD,
        is_fixed=False,
        is_active=True,
    )
    db.add(committee)
    db.flush()
    return committee


def _add_member(db: Session, committee: Committee, user: User) -> None:
    db.add(CommitteeMember(committee_id=committee.id, user_id=user.id, is_active=True))
    db.flush()


def _create_risk(db: Session, committee: Committee, creator: User) -> RiskRecord:
    risk = RiskRecord(
        risk_id="RISK-MINUTES-REPORT",
        problem_description="Committee minutes report risk",
        domain=RiskDomain.FLIGHT_TEST,
        board_of_origin_id=committee.id,
        workflow_status=RiskWorkflowStatus.SUBMITTED_TO_OPERATIONAL_BOARD,
        lifecycle_status=RiskLifecycleStatus.OPEN,
        created_by_user_id=creator.id,
        is_active=True,
    )
    db.add(risk)
    db.flush()
    return risk


def _create_decision(
    db: Session,
    committee: Committee,
    risk: RiskRecord,
    user: User,
) -> RiskDecision:
    decision = RiskDecision(
        risk_record_id=risk.id,
        committee_id=committee.id,
        decision_type=RiskDecisionType.APPROVE,
        decision_text="Decision Record text",
        decided_by_user_id=user.id,
        decided_at=datetime.now(timezone.utc),
    )
    db.add(decision)
    db.flush()
    return decision


def _create_meeting(db: Session, committee: Committee, user: User):
    risk = _create_risk(db, committee, user)
    decision = _create_decision(db, committee, risk, user)
    from app.services.committee_meeting_service import create_committee_meeting

    return create_committee_meeting(
        db,
        created_by_user_id=user.id,
        data=CommitteeMeetingCreate(
            committee_id=committee.id,
            title="Monthly Committee Meeting Minutes",
            meeting_date=datetime(2026, 7, 7).date(),
            location="SMS room",
            agenda_summary="Agenda Summary",
            discussion_summary="General Discussion Summary",
            decisions_summary="Decision Summary",
            action_items_summary="Action Items",
            attendees=[{"user_id": user.id, "role_label": "Chair"}],
            risk_items=[
                {
                    "risk_record_id": risk.id,
                    "agenda_item_number": 1,
                    "discussion_summary": "Discussion Summary",
                    "decision_summary": "Decision Summary",
                    "action_items": "Action Items",
                    "linked_risk_decision_id": decision.id,
                    "follow_up_required": True,
                    "follow_up_notes": "Follow-up notes",
                }
            ],
        ),
    )


def test_authorized_user_generates_tracked_minutes_docx_with_required_content(
    db_session: Session,
    tmp_path: Path,
) -> None:
    user = _create_user(db_session)
    committee = _create_committee(db_session)
    _add_member(db_session, committee, user)
    meeting = _create_meeting(db_session, committee, user)

    report = generate_and_track_committee_meeting_minutes_report(
        db_session,
        meeting_id=meeting.id,
        output_dir=tmp_path,
        generated_by_user_id=user.id,
    )
    text = _docx_text(Path(report.file_path))

    assert report.report_type == COMMITTEE_MEETING_MINUTES_REPORT_TYPE
    assert report.committee_id == committee.id
    assert report.risk_record_id is None
    for expected in (
        "Committee Meeting Minutes",
        committee.name,
        "Authority Level",
        "Attendance",
        "Risk Agenda Items",
        "Decision Summary",
        "Decision Record",
        "Status",
        "DRAFT",
    ):
        assert expected in text


def test_finalized_minutes_report_contains_finalized_status(
    db_session: Session,
    tmp_path: Path,
) -> None:
    user = _create_user(db_session)
    committee = _create_committee(db_session)
    _add_member(db_session, committee, user)
    meeting = _create_meeting(db_session, committee, user)
    finalized = finalize_committee_meeting(
        db_session,
        meeting_id=meeting.id,
        data=CommitteeMeetingFinalize(finalization_notes="Finalized"),
        finalized_by_user_id=user.id,
    )

    report = generate_and_track_committee_meeting_minutes_report(
        db_session,
        meeting_id=finalized.id,
        output_dir=tmp_path,
        generated_by_user_id=user.id,
    )
    text = _docx_text(Path(report.file_path))

    assert "FINALIZED" in text


def test_unauthorized_user_cannot_generate_minutes_report(
    db_session: Session,
    tmp_path: Path,
) -> None:
    member = _create_user(db_session)
    unrelated = _create_user(db_session)
    committee = _create_committee(db_session)
    _add_member(db_session, committee, member)
    meeting = _create_meeting(db_session, committee, member)

    with pytest.raises(ReportTrackingBusinessRuleError, match="not authorized|not an active"):
        generate_and_track_committee_meeting_minutes_report(
            db_session,
            meeting_id=meeting.id,
            output_dir=tmp_path,
            generated_by_user_id=unrelated.id,
        )


def test_minutes_report_api_generates_and_downloads_docx(
    client: TestClient,
    db_session: Session,
    tmp_path: Path,
) -> None:
    user = _create_user(db_session)
    committee = _create_committee(db_session)
    _add_member(db_session, committee, user)
    meeting = _create_meeting(db_session, committee, user)

    generated = client.post(
        f"/reports/committee-meeting-minutes/{meeting.id}",
        headers=_headers(user),
        json={"output_dir": str(tmp_path)},
    )
    downloaded = client.get(
        f"/reports/{generated.json()['id']}/download",
        headers=_headers(user),
    )

    assert generated.status_code == 201
    assert generated.json()["report_type"] == COMMITTEE_MEETING_MINUTES_REPORT_TYPE
    assert downloaded.status_code == 200
    assert downloaded.headers["content-type"] == (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    assert downloaded.content

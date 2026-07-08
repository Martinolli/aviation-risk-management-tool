import csv
import io
import uuid
from collections.abc import Generator
from datetime import datetime, timedelta, timezone
from pathlib import Path

import pytest
from docx import Document
from fastapi.testclient import TestClient
from sqlalchemy import create_engine
from sqlalchemy.orm import Session, sessionmaker
from sqlalchemy.pool import StaticPool

import app.models  # noqa: F401
from app.core.database import get_db
from app.main import app
from app.models.audit import AuditLog
from app.models.base import Base
from app.models.committee import Committee, CommitteeMember
from app.models.enums import (
    AuditAction,
    AuthorityLevel,
    CommitteeType,
    RiskDomain,
    RiskLifecycleStatus,
    RiskWorkflowStatus,
)
from app.models.risk import RiskRecord
from app.models.user import User
from app.services.auth_service import create_access_token


@pytest.fixture()
def db_session() -> Generator[Session, None, None]:
    engine = create_engine(
        "sqlite+pysqlite:///:memory:",
        connect_args={"check_same_thread": False},
        poolclass=StaticPool,
    )
    Base.metadata.create_all(engine)
    SessionLocal = sessionmaker(bind=engine, autoflush=False, autocommit=False)

    with SessionLocal() as session:
        yield session

    Base.metadata.drop_all(engine)


@pytest.fixture()
def client(
    db_session: Session,
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> Generator[TestClient, None, None]:
    monkeypatch.setattr("app.api.audit_logs.AUDIT_EXPORT_OUTPUT_DIR", tmp_path)

    def override_get_db() -> Generator[Session, None, None]:
        yield db_session

    app.dependency_overrides[get_db] = override_get_db
    with TestClient(app) as test_client:
        yield test_client
    app.dependency_overrides.clear()


def _create_user(db: Session, *, is_active: bool = True) -> User:
    user = User(
        email=f"audit-export-{uuid.uuid4()}@example.com",
        display_name="Audit Export User",
        is_active=is_active,
    )
    db.add(user)
    db.commit()
    db.refresh(user)
    return user


def _headers(user: User) -> dict[str, str]:
    return {"Authorization": f"Bearer {create_access_token(user_id=user.id)}"}


def _create_governance_reader(db: Session) -> User:
    user = _create_user(db)
    committee = Committee(
        name=f"Audit Export Governance {uuid.uuid4()}",
        authority_level=AuthorityLevel.MIDDLE,
        committee_type=CommitteeType.RISK_MANAGEMENT_COMMITTEE,
        is_fixed=True,
        is_active=True,
    )
    db.add(committee)
    db.commit()
    db.add(CommitteeMember(committee_id=committee.id, user_id=user.id, is_active=True))
    db.commit()
    return user


def _create_risk(db: Session, *, creator: User) -> RiskRecord:
    risk = RiskRecord(
        problem_description=f"Audit export risk {uuid.uuid4()}",
        domain=RiskDomain.FLIGHT_TEST,
        workflow_status=RiskWorkflowStatus.DRAFT,
        lifecycle_status=RiskLifecycleStatus.OPEN,
        created_by_user_id=creator.id,
        is_active=True,
    )
    db.add(risk)
    db.commit()
    db.refresh(risk)
    return risk


def _create_audit_log(
    db: Session,
    *,
    entity_type: str = "RiskRecord",
    entity_id: uuid.UUID | None = None,
    action: AuditAction = AuditAction.CREATE,
    changed_by_user_id: uuid.UUID | None = None,
    changed_at: datetime | None = None,
    field_name: str | None = "workflow_status",
    old_value: object = None,
    new_value: object = None,
    reason: str | None = "Traceability check",
) -> AuditLog:
    audit_log = AuditLog(
        entity_type=entity_type,
        entity_id=entity_id or uuid.uuid4(),
        action=action,
        field_name=field_name,
        old_value=old_value,
        new_value=new_value if new_value is not None else {"status": "recorded"},
        changed_by_user_id=changed_by_user_id,
        changed_at=changed_at or datetime.now(timezone.utc),
        reason=reason,
    )
    db.add(audit_log)
    db.commit()
    db.refresh(audit_log)
    return audit_log


def _csv_rows(response) -> list[dict[str, str]]:
    content = response.content.decode("utf-8")
    return list(csv.DictReader(io.StringIO(content)))


def _docx_text(response) -> str:
    document = Document(io.BytesIO(response.content))
    parts: list[str] = []
    parts.extend(paragraph.text for paragraph in document.paragraphs)
    for table in document.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


def test_authorized_user_can_export_csv(
    client: TestClient, db_session: Session
) -> None:
    reader = _create_governance_reader(db_session)
    _create_audit_log(db_session, entity_type="Committee")

    response = client.get("/audit-logs/export/csv", headers=_headers(reader))

    assert response.status_code == 200
    assert response.headers["content-type"].startswith("text/csv")


def test_csv_contains_expected_headers(
    client: TestClient, db_session: Session
) -> None:
    reader = _create_governance_reader(db_session)
    _create_audit_log(db_session, entity_type="Committee")

    response = client.get("/audit-logs/export/csv", headers=_headers(reader))
    header = response.content.decode("utf-8").splitlines()[0]

    assert header.split(",") == [
        "audit_log_id",
        "entity_type",
        "entity_id",
        "action",
        "field_name",
        "old_value_json",
        "new_value_json",
        "changed_by_user_id",
        "changed_at_utc",
        "reason",
        "created_at",
        "updated_at",
    ]


def test_csv_contains_authorized_audit_log_values(
    client: TestClient, db_session: Session
) -> None:
    reader = _create_governance_reader(db_session)
    audit_log = _create_audit_log(
        db_session,
        entity_type="Committee",
        action=AuditAction.UPDATE,
        field_name="name",
        old_value={"name": "Old"},
        new_value={"name": "New"},
    )

    response = client.get("/audit-logs/export/csv", headers=_headers(reader))
    rows = _csv_rows(response)

    assert rows[0]["audit_log_id"] == str(audit_log.id)
    assert rows[0]["entity_type"] == "Committee"
    assert rows[0]["action"] == "UPDATE"
    assert '"name": "New"' in rows[0]["new_value_json"]


def test_csv_does_not_contain_unauthorized_audit_logs(
    client: TestClient, db_session: Session
) -> None:
    reader = _create_user(db_session)
    other = _create_user(db_session)
    readable_risk = _create_risk(db_session, creator=reader)
    unreadable_risk = _create_risk(db_session, creator=other)
    readable_log = _create_audit_log(
        db_session, entity_type="RiskRecord", entity_id=readable_risk.id
    )
    unreadable_log = _create_audit_log(
        db_session, entity_type="RiskRecord", entity_id=unreadable_risk.id
    )

    response = client.get("/audit-logs/export/csv", headers=_headers(reader))
    exported_ids = {row["audit_log_id"] for row in _csv_rows(response)}

    assert str(readable_log.id) in exported_ids
    assert str(unreadable_log.id) not in exported_ids


def test_authorized_user_can_export_docx(
    client: TestClient, db_session: Session
) -> None:
    reader = _create_governance_reader(db_session)
    _create_audit_log(db_session, entity_type="Committee")

    response = client.get("/audit-logs/export/docx", headers=_headers(reader))

    assert response.status_code == 200
    assert response.headers["content-type"] == (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )


def test_docx_contains_controlled_export_content(
    client: TestClient, db_session: Session
) -> None:
    reader = _create_governance_reader(db_session)
    _create_audit_log(db_session, entity_type="Committee")

    response = client.get("/audit-logs/export/docx", headers=_headers(reader))
    text = _docx_text(response)

    for expected in [
        "Audit Trail Export",
        "SMS Risk Management Process Tool",
        "Controlled Export",
        "Entity Type",
        "Action",
        "Audit Trail Records",
        "Disclaimer",
    ]:
        assert expected in text


def test_entity_type_filter_works_in_csv_export(
    client: TestClient, db_session: Session
) -> None:
    reader = _create_governance_reader(db_session)
    target = _create_audit_log(db_session, entity_type="RiskRecord")
    _create_audit_log(db_session, entity_type="Committee")

    response = client.get(
        "/audit-logs/export/csv?entity_type=RiskRecord",
        headers=_headers(reader),
    )

    assert [row["audit_log_id"] for row in _csv_rows(response)] == [str(target.id)]


def test_entity_id_filter_works_in_csv_export(
    client: TestClient, db_session: Session
) -> None:
    reader = _create_governance_reader(db_session)
    entity_id = uuid.uuid4()
    target = _create_audit_log(db_session, entity_id=entity_id)
    _create_audit_log(db_session)

    response = client.get(
        f"/audit-logs/export/csv?entity_id={entity_id}",
        headers=_headers(reader),
    )

    assert [row["audit_log_id"] for row in _csv_rows(response)] == [str(target.id)]


def test_action_filter_works_in_csv_export(
    client: TestClient, db_session: Session
) -> None:
    reader = _create_governance_reader(db_session)
    target = _create_audit_log(db_session, action=AuditAction.GENERATE_REPORT)
    _create_audit_log(db_session, action=AuditAction.CREATE)

    response = client.get(
        "/audit-logs/export/csv?action=GENERATE_REPORT",
        headers=_headers(reader),
    )

    assert [row["audit_log_id"] for row in _csv_rows(response)] == [str(target.id)]


def test_changed_by_filter_works_in_csv_export(
    client: TestClient, db_session: Session
) -> None:
    reader = _create_governance_reader(db_session)
    changed_by = _create_user(db_session)
    target = _create_audit_log(db_session, changed_by_user_id=changed_by.id)
    _create_audit_log(db_session, changed_by_user_id=uuid.uuid4())

    response = client.get(
        f"/audit-logs/export/csv?changed_by_user_id={changed_by.id}",
        headers=_headers(reader),
    )

    assert [row["audit_log_id"] for row in _csv_rows(response)] == [str(target.id)]


def test_changed_at_from_filter_works(
    client: TestClient, db_session: Session
) -> None:
    reader = _create_governance_reader(db_session)
    now = datetime.now(timezone.utc)
    _create_audit_log(db_session, changed_at=now - timedelta(days=2))
    target = _create_audit_log(db_session, changed_at=now)

    response = client.get(
        "/audit-logs/export/csv",
        params={"changed_at_from": (now - timedelta(hours=1)).isoformat()},
        headers=_headers(reader),
    )

    assert [row["audit_log_id"] for row in _csv_rows(response)] == [str(target.id)]


def test_changed_at_to_filter_works(
    client: TestClient, db_session: Session
) -> None:
    reader = _create_governance_reader(db_session)
    now = datetime.now(timezone.utc)
    target = _create_audit_log(db_session, changed_at=now - timedelta(days=2))
    _create_audit_log(db_session, changed_at=now)

    response = client.get(
        "/audit-logs/export/csv",
        params={"changed_at_to": (now - timedelta(days=1)).isoformat()},
        headers=_headers(reader),
    )

    assert [row["audit_log_id"] for row in _csv_rows(response)] == [str(target.id)]


def test_invalid_date_range_returns_http_400(
    client: TestClient, db_session: Session
) -> None:
    reader = _create_governance_reader(db_session)
    now = datetime.now(timezone.utc)

    response = client.get(
        "/audit-logs/export/csv",
        params={
            "changed_at_from": now.isoformat(),
            "changed_at_to": (now - timedelta(days=1)).isoformat(),
        },
        headers=_headers(reader),
    )

    assert response.status_code == 400


def test_limit_below_one_returns_http_400(
    client: TestClient, db_session: Session
) -> None:
    reader = _create_governance_reader(db_session)

    response = client.get("/audit-logs/export/csv?limit=0", headers=_headers(reader))

    assert response.status_code == 400


def test_excessive_limit_returns_http_400(
    client: TestClient, db_session: Session
) -> None:
    reader = _create_governance_reader(db_session)

    response = client.get(
        "/audit-logs/export/csv?limit=5001",
        headers=_headers(reader),
    )

    assert response.status_code == 400


def test_unauthenticated_user_cannot_export(client: TestClient) -> None:
    response = client.get("/audit-logs/export/csv")

    assert response.status_code == 400


def test_inactive_user_cannot_export(
    client: TestClient, db_session: Session
) -> None:
    inactive = _create_user(db_session, is_active=False)

    response = client.get(
        "/audit-logs/export/csv",
        headers={"X-User-Id": str(inactive.id)},
    )

    assert response.status_code == 403


def test_existing_audit_log_list_behavior_still_works(
    client: TestClient, db_session: Session
) -> None:
    reader = _create_governance_reader(db_session)
    _create_audit_log(db_session, entity_type="Committee")

    response = client.get("/audit-logs", headers=_headers(reader))

    assert response.status_code == 200
    assert len(response.json()) == 1


def test_existing_audit_log_get_behavior_still_works(
    client: TestClient, db_session: Session
) -> None:
    reader = _create_governance_reader(db_session)
    audit_log = _create_audit_log(db_session, entity_type="Committee")

    response = client.get(f"/audit-logs/{audit_log.id}", headers=_headers(reader))

    assert response.status_code == 200
    assert response.json()["id"] == str(audit_log.id)


def test_export_files_are_created_in_tmp_path(
    client: TestClient, db_session: Session, tmp_path: Path
) -> None:
    reader = _create_governance_reader(db_session)
    _create_audit_log(db_session, entity_type="Committee")

    csv_response = client.get("/audit-logs/export/csv", headers=_headers(reader))
    docx_response = client.get("/audit-logs/export/docx", headers=_headers(reader))

    assert csv_response.status_code == 200
    assert docx_response.status_code == 200
    assert list(tmp_path.glob("audit_trail_export_*.csv"))
    assert list(tmp_path.glob("audit_trail_export_*.docx"))

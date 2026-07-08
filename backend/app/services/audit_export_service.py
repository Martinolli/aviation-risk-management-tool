import csv
import json
import uuid
from datetime import date, datetime, timezone
from enum import Enum
from pathlib import Path
from typing import Any

from docx import Document
from docx.document import Document as DocumentObject
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from sqlalchemy.orm import Session

from app.models.audit import AuditLog
from app.schemas.audit_export import AuditLogExportFilters
from app.services.audit_query_service import (
    AuditQueryBusinessRuleError,
    list_audit_logs,
)

DEFAULT_AUDIT_EXPORT_OUTPUT_DIR = Path("generated_reports")
AUDIT_DOCX_VALUE_MAX_LENGTH = 350
NAVY = "17365D"
LIGHT_BLUE = "DCE6F1"


class AuditExportBusinessRuleError(ValueError):
    pass


def _utc_now() -> datetime:
    return datetime.now(timezone.utc)


def _timestamp_for_filename(value: datetime) -> str:
    return value.strftime("%Y%m%d_%H%M%S")


def _ensure_output_dir(output_dir: Path | str) -> Path:
    output_path = Path(output_dir)
    output_path.mkdir(parents=True, exist_ok=True)
    return output_path


def _json_default(value: Any) -> str:
    if isinstance(value, datetime):
        return value.isoformat()
    if isinstance(value, date):
        return value.isoformat()
    if isinstance(value, uuid.UUID):
        return str(value)
    if isinstance(value, Enum):
        return value.value
    return str(value)


def _json_text(value: Any) -> str:
    return json.dumps(value, default=_json_default, ensure_ascii=False)


def _safe_text(value: Any) -> str:
    if value is None:
        return "Not recorded"
    if isinstance(value, Enum):
        return value.value
    if isinstance(value, uuid.UUID):
        return str(value)
    if isinstance(value, datetime):
        utc_value = value.astimezone(timezone.utc) if value.tzinfo else value
        return f"{utc_value.strftime('%Y-%m-%d %H:%M:%S')} UTC"
    if isinstance(value, date):
        return value.isoformat()
    text = str(value)
    return text if text else "Not recorded"


def _compact_value(value: Any) -> str:
    if value is None:
        return "Not recorded"
    if isinstance(value, dict | list):
        text = json.dumps(
            value,
            default=_json_default,
            ensure_ascii=True,
            separators=(", ", ": "),
            sort_keys=isinstance(value, dict),
        )
    else:
        text = _safe_text(value)
    if len(text) <= AUDIT_DOCX_VALUE_MAX_LENGTH:
        return text
    return f"{text[: AUDIT_DOCX_VALUE_MAX_LENGTH - 3]}..."


def _format_filter(value: Any) -> str:
    return "All" if value is None else _safe_text(value)


def get_authorized_audit_logs_for_export(
    db: Session,
    *,
    requested_by_user_id: uuid.UUID | None,
    filters: AuditLogExportFilters,
) -> list[AuditLog]:
    try:
        return list_audit_logs(
            db,
            requested_by_user_id=requested_by_user_id,
            entity_type=filters.entity_type,
            entity_id=filters.entity_id,
            action=filters.action,
            changed_by_user_id=filters.changed_by_user_id,
            changed_at_from=filters.changed_at_from,
            changed_at_to=filters.changed_at_to,
            limit=filters.limit,
            offset=filters.offset,
            max_limit=5000,
        )
    except AuditQueryBusinessRuleError as exc:
        raise AuditExportBusinessRuleError(str(exc)) from exc


def export_audit_logs_csv(
    db: Session,
    *,
    requested_by_user_id: uuid.UUID | None,
    filters: AuditLogExportFilters,
    output_dir: Path | str = DEFAULT_AUDIT_EXPORT_OUTPUT_DIR,
) -> Path:
    audit_logs = get_authorized_audit_logs_for_export(
        db,
        requested_by_user_id=requested_by_user_id,
        filters=filters,
    )
    generated_at = _utc_now()
    file_path = (
        _ensure_output_dir(output_dir)
        / f"audit_trail_export_{_timestamp_for_filename(generated_at)}.csv"
    )

    with file_path.open("w", encoding="utf-8", newline="") as csv_file:
        writer = csv.DictWriter(
            csv_file,
            fieldnames=[
                "audit_log_id",
                "entity_type",
                "entity_id",
                "action",
                "field_name",
                "old_value_json",
                "new_value_json",
                "changed_by_user_id",
                "changed_at_utc",
                "reason",
                "created_at",
                "updated_at",
            ],
        )
        writer.writeheader()
        for audit_log in audit_logs:
            writer.writerow(
                {
                    "audit_log_id": str(audit_log.id),
                    "entity_type": audit_log.entity_type,
                    "entity_id": str(audit_log.entity_id),
                    "action": audit_log.action.value,
                    "field_name": audit_log.field_name or "",
                    "old_value_json": _json_text(audit_log.old_value),
                    "new_value_json": _json_text(audit_log.new_value),
                    "changed_by_user_id": (
                        str(audit_log.changed_by_user_id)
                        if audit_log.changed_by_user_id is not None
                        else ""
                    ),
                    "changed_at_utc": _safe_text(audit_log.changed_at),
                    "reason": audit_log.reason or "",
                    "created_at": _safe_text(audit_log.created_at),
                    "updated_at": _safe_text(audit_log.updated_at),
                }
            )

    return file_path


def export_audit_logs_docx(
    db: Session,
    *,
    requested_by_user_id: uuid.UUID | None,
    filters: AuditLogExportFilters,
    output_dir: Path | str = DEFAULT_AUDIT_EXPORT_OUTPUT_DIR,
) -> Path:
    audit_logs = get_authorized_audit_logs_for_export(
        db,
        requested_by_user_id=requested_by_user_id,
        filters=filters,
    )
    generated_at = _utc_now()
    file_path = (
        _ensure_output_dir(output_dir)
        / f"audit_trail_export_{_timestamp_for_filename(generated_at)}.docx"
    )

    document = Document()
    _set_document_defaults(document)
    _add_header(
        document,
        generated_at=generated_at,
        generated_by_user_id=requested_by_user_id,
        filters=filters,
        record_count=len(audit_logs),
    )
    document.add_heading("Export Scope", level=1)
    document.add_paragraph(
        "This Audit Trail Export contains audit records authorized for the "
        "requesting user and filtered according to the export criteria."
    )
    document.add_heading("Audit Trail Records", level=1)
    _add_audit_log_table(document, audit_logs)
    document.add_heading("Disclaimer", level=1)
    document.add_paragraph(
        "This Audit Trail Export is generated from system records. It supports "
        "SMS governance traceability, audit preparation, investigation review, "
        "and compliance evidence. It does not replace formal committee "
        "decisions, accountable manager acceptance, or required regulatory "
        "documentation."
    )

    document.save(file_path)
    return file_path


def _set_document_defaults(document: DocumentObject) -> None:
    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)

    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(9)
    styles["Title"].font.name = "Arial"
    styles["Title"].font.size = Pt(24)
    styles["Title"].font.color.rgb = RGBColor(23, 54, 93)
    for style_name, font_size in (("Heading 1", 14), ("Heading 2", 11)):
        style = styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(font_size)
        style.font.color.rgb = RGBColor(23, 54, 93)


def _add_header(
    document: DocumentObject,
    *,
    generated_at: datetime,
    generated_by_user_id: uuid.UUID | None,
    filters: AuditLogExportFilters,
    record_count: int,
) -> None:
    title = document.add_heading("Audit Trail Export", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run("SMS Risk Management Process Tool")
    subtitle_run.bold = True
    subtitle_run.font.size = Pt(12)
    subtitle_run.font.color.rgb = RGBColor(72, 101, 129)

    metadata_values = [
        ("Export Type", "Controlled Export"),
        ("Generated At UTC", _safe_text(generated_at)),
        ("Generated By User ID", _format_filter(generated_by_user_id)),
        ("Entity Type filter", _format_filter(filters.entity_type)),
        ("Entity ID filter", _format_filter(filters.entity_id)),
        ("Action filter", _format_filter(filters.action)),
        ("Changed By filter", _format_filter(filters.changed_by_user_id)),
        ("Changed At From", _format_filter(filters.changed_at_from)),
        ("Changed At To", _format_filter(filters.changed_at_to)),
        ("Record Count", str(record_count)),
    ]
    metadata = document.add_table(rows=5, cols=2)
    metadata.style = "Table Grid"
    metadata.alignment = WD_TABLE_ALIGNMENT.CENTER

    for cell, (label, value) in zip(
        (cell for row in metadata.rows for cell in row.cells),
        metadata_values,
        strict=True,
    ):
        cell.text = ""
        paragraph = cell.paragraphs[0]
        paragraph.add_run(f"{label}: ").bold = True
        paragraph.add_run(value)

    footer = document.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.text = "Audit Trail Export - Controlled Export"


def _add_audit_log_table(
    document: DocumentObject,
    audit_logs: list[AuditLog],
) -> None:
    if not audit_logs:
        document.add_paragraph("No authorized audit records matched the export criteria.")
        return

    headers = [
        "Changed At UTC",
        "Entity Type",
        "Entity ID",
        "Action",
        "Field",
        "Changed By",
        "Reason",
        "Old Value",
        "New Value",
    ]
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for cell, header in zip(table.rows[0].cells, headers, strict=True):
        cell.text = header
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        shading = cell._tc.get_or_add_tcPr()
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        fill = OxmlElement("w:shd")
        fill.set(qn("w:fill"), LIGHT_BLUE)
        shading.append(fill)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(23, 54, 93)
                run.font.size = Pt(8)

    for audit_log in audit_logs:
        values = [
            _safe_text(audit_log.changed_at),
            audit_log.entity_type,
            str(audit_log.entity_id),
            audit_log.action.value,
            audit_log.field_name or "",
            _format_filter(audit_log.changed_by_user_id),
            audit_log.reason or "",
            _compact_value(audit_log.old_value),
            _compact_value(audit_log.new_value),
        ]
        row = table.add_row()
        for cell, value in zip(row.cells, values, strict=True):
            cell.text = value
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(7)

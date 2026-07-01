import json
import re
import uuid
from datetime import date, datetime, timezone
from enum import Enum
from pathlib import Path
from typing import Any

from docx import Document
from docx.document import Document as DocumentObject
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from sqlalchemy import select
from sqlalchemy.orm import Session

from app.models.audit import AuditLog
from app.models.committee import Committee
from app.models.risk import (
    RiskAction,
    RiskAssessment,
    RiskDecision,
    RiskEvidence,
    RiskRecord,
)
from app.services.risk_detail_service import _build_risk_record_detail

REPORT_TITLE = "Risk Dossier Report"
REPORT_SUBTITLE = "SMS Risk Management Process Tool"
AUDIT_VALUE_MAX_LENGTH = 400
NAVY = "17365D"
LIGHT_BLUE = "DCE6F1"


class ReportGenerationError(ValueError):
    pass


class ReportRiskNotFoundError(ReportGenerationError):
    pass


def _format_label(value: Any) -> str:
    if isinstance(value, Enum):
        value = value.value
    text = str(value)
    if re.fullmatch(r"[A-Z0-9_]+", text):
        return text.replace("_", " ").title()
    return text


def _format_datetime(value: Any) -> str:
    if isinstance(value, datetime):
        utc_value = value.astimezone(timezone.utc) if value.tzinfo else value
        return f"{utc_value.strftime('%Y-%m-%d %H:%M:%S')} UTC"
    if isinstance(value, date):
        return value.isoformat()
    return _safe_text(value)


def _format_bool(value: bool | None) -> str:
    if value is True:
        return "Yes"
    if value is False:
        return "No"
    return "Not recorded."


def _format_file_size(file_size_bytes: int) -> str:
    if file_size_bytes < 1024:
        return f"{file_size_bytes} B"
    if file_size_bytes < 1024 * 1024:
        return f"{file_size_bytes / 1024:.1f} KB"
    return f"{file_size_bytes / (1024 * 1024):.1f} MB"


def _safe_text(value: Any) -> str:
    if value is None:
        return "Not recorded."
    if isinstance(value, Enum):
        return _format_label(value)
    if isinstance(value, datetime | date):
        return value.isoformat()
    if isinstance(value, uuid.UUID):
        return str(value)
    text = str(value)
    return text if text else "Not recorded."


def _compact_audit_value(value: Any) -> str:
    if value is None:
        return "Not recorded."
    if isinstance(value, dict | list):
        text = json.dumps(
            value,
            default=_safe_text,
            ensure_ascii=True,
            separators=(", ", ": "),
            sort_keys=isinstance(value, dict),
        )
    else:
        text = _safe_text(value)
    if len(text) <= AUDIT_VALUE_MAX_LENGTH:
        return text
    return f"{text[: AUDIT_VALUE_MAX_LENGTH - 3]}..."


def _set_document_defaults(document: DocumentObject) -> None:
    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)

    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(9.5)
    styles["Title"].font.name = "Arial"
    styles["Title"].font.size = Pt(25)
    styles["Title"].font.color.rgb = RGBColor(23, 54, 93)
    for style_name, font_size in (("Heading 1", 15), ("Heading 2", 11)):
        style = styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(font_size)
        style.font.color.rgb = RGBColor(23, 54, 93)


def _add_report_header(
    document: DocumentObject,
    risk_record: RiskRecord,
    *,
    generated_at: datetime,
) -> None:
    title = document.add_heading(REPORT_TITLE, 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run(REPORT_SUBTITLE)
    subtitle_run.bold = True
    subtitle_run.font.size = Pt(12)
    subtitle_run.font.color.rgb = RGBColor(72, 101, 129)

    metadata = document.add_table(rows=2, cols=3)
    metadata.style = "Table Grid"
    metadata.alignment = WD_TABLE_ALIGNMENT.CENTER
    values = (
        ("Risk ID", risk_record.risk_id),
        ("Generated at UTC", _format_datetime(generated_at)),
        ("Workflow Status", _format_label(risk_record.workflow_status)),
        ("Lifecycle Status", _format_label(risk_record.lifecycle_status)),
        ("Internal UUID", risk_record.id),
        ("Report Type", REPORT_TITLE),
    )
    for cell, (label, value) in zip(
        (cell for row in metadata.rows for cell in row.cells),
        values,
        strict=True,
    ):
        cell.text = ""
        paragraph = cell.paragraphs[0]
        paragraph.add_run(f"{label}: ").bold = True
        paragraph.add_run(_safe_text(value))

    scope = document.add_paragraph()
    scope.add_run("Report Scope: ").bold = True
    scope.add_run(
        "This dossier consolidates the risk record, risk package, assessments, "
        "mitigation actions, committee decisions, evidence metadata, and audit "
        "trail available in the Risk Management Process Tool at the time of "
        "generation."
    )

    footer = document.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run(f"{REPORT_TITLE} | {_safe_text(risk_record.risk_id)}")


def _add_section_heading(document: DocumentObject, title: str) -> None:
    heading = document.add_heading(title, level=1)
    heading.paragraph_format.keep_with_next = True


def _add_key_value_paragraph(
    document: DocumentObject,
    key: str,
    value: Any,
) -> None:
    paragraph = document.add_paragraph()
    paragraph.add_run(f"{key}: ").bold = True
    paragraph.add_run(_safe_text(value))


def _add_bullet_list(document: DocumentObject, items: list[str] | None) -> None:
    if not items:
        document.add_paragraph("Not recorded.")
        return
    for item in items:
        document.add_paragraph(_safe_text(item), style="List Bullet")


def _shade_cell(cell, fill: str) -> None:
    cell_properties = cell._tc.get_or_add_tcPr()
    shading = cell_properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        cell_properties.append(shading)
    shading.set(qn("w:fill"), fill)


def _repeat_header_row(row) -> None:
    table_row_properties = row._tr.get_or_add_trPr()
    table_header = OxmlElement("w:tblHeader")
    table_header.set(qn("w:val"), "true")
    table_row_properties.append(table_header)


def _add_table(
    document: DocumentObject,
    headers: list[str],
    rows: list[list[Any]],
    *,
    style: str = "Table Grid",
) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = style
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    header_cells = table.rows[0].cells
    _repeat_header_row(table.rows[0])
    for index, header in enumerate(headers):
        header_cells[index].text = header
        header_cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        _shade_cell(header_cells[index], NAVY)
        for run in header_cells[index].paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.size = Pt(8)

    for row_number, row in enumerate(rows, start=1):
        row_cells = table.add_row().cells
        for index, value in enumerate(row):
            row_cells[index].text = _safe_text(value)
            row_cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            if row_number % 2 == 0:
                _shade_cell(row_cells[index], LIGHT_BLUE)
            for paragraph in row_cells[index].paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.size = Pt(7.5)


def _safe_filename(value: str) -> str:
    filename = re.sub(r"[^A-Za-z0-9_.-]+", "_", value.strip())
    filename = filename.strip("._")
    return filename or "risk_dossier"


def _assessment_rows(assessments: list[RiskAssessment]) -> list[list[Any]]:
    return [
        [
            _format_label(assessment.assessment_type),
            assessment.severity,
            assessment.likelihood,
            assessment.risk_level,
            assessment.calculated_score,
            _format_bool(assessment.is_tolerable),
            _format_bool(assessment.requires_mitigation),
            _format_bool(assessment.requires_escalation),
            assessment.rationale,
            _format_datetime(assessment.assessed_at),
        ]
        for assessment in assessments
    ]


def _action_rows(actions: list[RiskAction]) -> list[list[Any]]:
    return [
        [
            action.title,
            _format_label(action.status),
            action.action_owner_user_id,
            _format_datetime(action.due_date),
            _format_datetime(action.completed_at),
            action.completion_notes,
        ]
        for action in actions
    ]


def _decision_rows(
    decisions: list[RiskDecision],
    committees: dict[uuid.UUID, Committee],
) -> list[list[Any]]:
    rows: list[list[Any]] = []
    for decision in decisions:
        committee = committees.get(decision.committee_id)
        rows.append(
            [
                _format_label(decision.decision_type),
                committee.name if committee is not None else decision.committee_id,
                (
                    _format_label(committee.authority_level)
                    if committee is not None
                    else None
                ),
                decision.decision_text,
                decision.decided_by_user_id,
                _format_datetime(decision.decided_at),
            ]
        )
    return rows


def _evidence_rows(evidence_items: list[RiskEvidence]) -> list[list[Any]]:
    return [
        [
            evidence.original_filename,
            evidence.description,
            evidence.content_type,
            _format_file_size(evidence.file_size_bytes),
            evidence.uploaded_by_user_id,
            _format_datetime(evidence.uploaded_at),
            "Active" if evidence.is_active else "Archived",
            _format_datetime(evidence.archived_at),
            evidence.archive_reason,
        ]
        for evidence in evidence_items
    ]


def _audit_rows(audit_logs: list[AuditLog]) -> list[list[Any]]:
    return [
        [
            _format_datetime(audit_log.changed_at),
            audit_log.entity_type,
            audit_log.entity_id,
            _format_label(audit_log.action),
            audit_log.field_name,
            audit_log.changed_by_user_id,
            audit_log.reason,
            _compact_audit_value(audit_log.old_value),
            _compact_audit_value(audit_log.new_value),
        ]
        for audit_log in audit_logs
    ]


def generate_risk_dossier_docx(
    db: Session,
    *,
    risk_record_id: uuid.UUID,
    output_dir: Path | str,
) -> Path:
    detail = _build_risk_record_detail(db, risk_record_id=risk_record_id)
    if detail is None:
        raise ReportRiskNotFoundError("Risk record not found")

    risk_record: RiskRecord = detail["risk_record"]
    assessments: list[RiskAssessment] = detail["assessments"]
    actions: list[RiskAction] = detail["actions"]
    decisions: list[RiskDecision] = detail["decisions"]
    evidence_items: list[RiskEvidence] = detail["evidence_items"]
    audit_logs: list[AuditLog] = detail["audit_logs"]
    audit_summary: dict[str, Any] = detail["audit_summary"]
    board_of_origin = (
        db.get(Committee, risk_record.board_of_origin_id)
        if risk_record.board_of_origin_id is not None
        else None
    )
    committee_ids = {decision.committee_id for decision in decisions}
    committees = {
        committee.id: committee
        for committee in db.scalars(
            select(Committee).where(Committee.id.in_(committee_ids))
        ).all()
    }

    output_path = Path(output_dir)
    output_path.mkdir(parents=True, exist_ok=True)
    base_name = risk_record.risk_id or f"risk_{risk_record.id}"
    file_path = output_path / f"{_safe_filename(base_name)}_dossier.docx"
    generated_at = datetime.now(timezone.utc)

    document = Document()
    _set_document_defaults(document)
    _add_report_header(document, risk_record, generated_at=generated_at)

    _add_section_heading(document, "Section 1 - Risk Record Summary")
    _add_key_value_paragraph(document, "Problem Description", risk_record.problem_description)
    _add_key_value_paragraph(document, "Source Trigger", risk_record.source_trigger)
    _add_key_value_paragraph(document, "Domain", _format_label(risk_record.domain))
    _add_key_value_paragraph(document, "System Scope", risk_record.system_scope)
    _add_key_value_paragraph(document, "Central Event", risk_record.central_event)
    _add_key_value_paragraph(document, "Hazard Statement", risk_record.hazard_statement)
    _add_key_value_paragraph(
        document, "Workflow Status", _format_label(risk_record.workflow_status)
    )
    _add_key_value_paragraph(
        document, "Lifecycle Status", _format_label(risk_record.lifecycle_status)
    )
    _add_key_value_paragraph(
        document,
        "Board of Origin / Originating Committee",
        board_of_origin.name if board_of_origin is not None else None,
    )
    _add_key_value_paragraph(document, "Board of Origin ID", risk_record.board_of_origin_id)
    _add_key_value_paragraph(
        document,
        "Board of Origin Authority Level",
        (
            _format_label(board_of_origin.authority_level)
            if board_of_origin is not None
            else None
        ),
    )
    _add_key_value_paragraph(document, "Owner User ID", risk_record.owner_user_id)
    _add_key_value_paragraph(
        document, "Created At", _format_datetime(risk_record.created_at)
    )
    _add_key_value_paragraph(
        document, "Updated At", _format_datetime(risk_record.updated_at)
    )

    _add_section_heading(
        document, "Section 2 - Causes, Consequences, and Existing Controls"
    )
    document.add_heading("Causes", level=2)
    _add_bullet_list(document, risk_record.causes)
    document.add_heading("Consequences", level=2)
    _add_bullet_list(document, risk_record.consequences)
    document.add_heading("Existing Controls", level=2)
    _add_bullet_list(document, risk_record.existing_controls)

    _add_section_heading(document, "Section 3 - Risk Assessments")
    if assessments:
        _add_table(
            document,
            [
                "Type",
                "Severity",
                "Likelihood",
                "Risk Level",
                "Score",
                "Tolerable",
                "Mitigation",
                "Escalation",
                "Rationale",
                "Assessed At",
            ],
            _assessment_rows(assessments),
        )
    else:
        document.add_paragraph("No risk assessments recorded.")

    _add_section_heading(document, "Section 4 - Mitigation / Risk Actions")
    if actions:
        _add_table(
            document,
            [
                "Title",
                "Status",
                "Owner User ID",
                "Due Date",
                "Completed At",
                "Completion Notes",
            ],
            _action_rows(actions),
        )
    else:
        document.add_paragraph("No risk actions recorded.")

    _add_section_heading(document, "Section 5 - Committee Decisions")
    if decisions:
        _add_table(
            document,
            [
                "Decision Type",
                "Committee",
                "Authority Level",
                "Decision Text",
                "Decided By",
                "Decided At",
            ],
            _decision_rows(decisions, committees),
        )
    else:
        document.add_paragraph("No committee decisions recorded.")

    _add_section_heading(document, "Section 6 - Evidence / Supporting Documents")
    if evidence_items:
        _add_table(
            document,
            [
                "File Name",
                "Description",
                "Content Type",
                "File Size",
                "Uploaded By",
                "Uploaded At",
                "Status",
                "Archived At",
                "Archive Reason",
            ],
            _evidence_rows(evidence_items),
        )
    else:
        document.add_paragraph("No evidence or supporting documents recorded.")
    document.add_paragraph(
        "Evidence files are stored in the system repository/storage and are "
        "referenced by metadata in this report."
    )

    document.add_page_break()
    _add_section_heading(document, "Section 7 - Audit Summary")
    _add_key_value_paragraph(document, "Total Audit Records", audit_summary["total_count"])
    _add_key_value_paragraph(document, "Create Records", audit_summary["create_count"])
    _add_key_value_paragraph(document, "Update Records", audit_summary["update_count"])
    _add_key_value_paragraph(document, "Workflow Records", audit_summary["workflow_count"])
    _add_key_value_paragraph(document, "Evidence Records", audit_summary["evidence_count"])
    _add_key_value_paragraph(
        document,
        "Latest Audit Change",
        _format_datetime(audit_summary["latest_changed_at"]),
    )

    _add_section_heading(document, "Section 8 - Audit Trail Annex")
    if audit_logs:
        _add_table(
            document,
            [
                "Changed At",
                "Entity Type",
                "Entity ID",
                "Action",
                "Field",
                "Changed By",
                "Reason",
                "Old Value",
                "New Value",
            ],
            _audit_rows(audit_logs),
        )
    else:
        document.add_paragraph("No audit trail records available.")

    _add_section_heading(document, "Section 9 - Notes / Disclaimer")
    document.add_paragraph(
        "This report is generated from system records. It does not replace formal "
        "committee approval, accountable manager acceptance, or required regulatory "
        "documentation."
    )

    try:
        document.save(file_path)
    except Exception as exc:
        raise ReportGenerationError("Failed to generate risk dossier DOCX") from exc

    return file_path

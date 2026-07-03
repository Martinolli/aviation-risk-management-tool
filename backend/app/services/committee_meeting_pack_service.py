import uuid
from collections import Counter, defaultdict
from datetime import date, datetime, timezone
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
from sqlalchemy import select
from sqlalchemy.orm import Session

from app.models.committee import Committee
from app.models.enums import (
    RiskAssessmentType,
    RiskMonitoringStatus,
)
from app.models.risk import (
    RiskAction,
    RiskAssessment,
    RiskEvidence,
    RiskMonitoringReview,
    RiskRecord,
)
from app.models.user import User
from app.services.decision_queue_service import (
    DecisionQueueBusinessRuleError,
    get_decision_queue_for_committee,
    get_decision_queue_scope,
)
from app.services.report_service import (
    _add_key_value_paragraph,
    _add_section_heading,
    _add_table,
    _format_datetime,
    _format_label,
    _safe_filename,
    _safe_text,
    _set_document_defaults,
)
from app.services.risk_access_service import (
    RiskAccessBusinessRuleError,
    is_active_committee_member,
    validate_active_user,
)
from app.services.risk_action_service import (
    get_risk_action_due_status,
    is_action_open_for_alerts,
)

REPORT_TITLE = "Committee Meeting Pack"
REPORT_SUBTITLE = "SMS Risk Management Process Tool"
OPEN_MONITORING_STATUSES = {
    RiskMonitoringStatus.ACTIVE,
    RiskMonitoringStatus.DUE,
    RiskMonitoringStatus.OVERDUE,
}
MONITORING_STATUS_PRIORITY = {
    RiskMonitoringStatus.OVERDUE: 0,
    RiskMonitoringStatus.DUE: 1,
    RiskMonitoringStatus.ACTIVE: 2,
}
MEETING_AGENDA_ITEMS = [
    "Opening / quorum confirmation",
    "Review of risk decision queue",
    "Risk-by-risk discussion",
    "Mitigation / action follow-up",
    "Escalation / acceptance decisions",
    "Monitoring and review-cycle items",
    "Summary of decisions and action items",
    "Closure",
]


class CommitteeMeetingPackBusinessRuleError(ValueError):
    pass


def _validate_generator(
    db: Session,
    *,
    committee_id: uuid.UUID,
    generated_by_user_id: uuid.UUID,
) -> tuple[Committee, User]:
    try:
        user = validate_active_user(
            db,
            user_id=generated_by_user_id,
            context="Committee Meeting Pack generation",
        )
    except RiskAccessBusinessRuleError as exc:
        raise CommitteeMeetingPackBusinessRuleError(str(exc)) from exc

    committee = db.get(Committee, committee_id)
    if committee is None:
        raise CommitteeMeetingPackBusinessRuleError("Committee does not exist")
    if not committee.is_active:
        raise CommitteeMeetingPackBusinessRuleError("Committee is inactive")
    if not is_active_committee_member(
        db,
        committee_id=committee.id,
        user_id=user.id,
    ):
        raise CommitteeMeetingPackBusinessRuleError(
            "User is not authorized to generate a meeting pack for this committee"
        )
    return committee, user


def _group_by_risk(items: list) -> dict[uuid.UUID, list]:
    grouped: dict[uuid.UUID, list] = defaultdict(list)
    for item in items:
        grouped[item.risk_record_id].append(item)
    return grouped


def _latest_assessments(
    assessments: list[RiskAssessment],
) -> dict[tuple[uuid.UUID, RiskAssessmentType], RiskAssessment]:
    latest: dict[tuple[uuid.UUID, RiskAssessmentType], RiskAssessment] = {}
    for assessment in sorted(
        assessments,
        key=lambda item: item.assessed_at,
        reverse=True,
    ):
        latest.setdefault(
            (assessment.risk_record_id, assessment.assessment_type),
            assessment,
        )
    return latest


def _assessment_summary(assessment: RiskAssessment | None) -> str:
    if assessment is None:
        return "Not recorded."
    return (
        f"Severity: {_safe_text(assessment.severity)}; "
        f"Likelihood: {_safe_text(assessment.likelihood)}; "
        f"Risk Level: {_safe_text(assessment.risk_level)}"
    )


def _monitoring_summary(reviews: list[RiskMonitoringReview]) -> str:
    active_reviews = [
        review
        for review in reviews
        if review.is_active and review.status in OPEN_MONITORING_STATUSES
    ]
    if not active_reviews:
        return "None"
    return _format_label(
        min(active_reviews, key=lambda item: MONITORING_STATUS_PRIORITY[item.status]).status
    )


def _risk_display_id(risk: RiskRecord) -> str:
    return risk.risk_id or str(risk.id)


def _queue_scope_text(scope: str | list[str]) -> str:
    return ", ".join(scope) if isinstance(scope, list) else scope


def _add_header(
    document,
    *,
    committee: Committee,
    generated_by_user_id: uuid.UUID,
    generated_at: datetime,
    queue_scope: str | list[str],
    meeting_title: str | None,
    meeting_date: date | None,
) -> None:
    title = document.add_heading(REPORT_TITLE, 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run(REPORT_SUBTITLE)
    subtitle_run.bold = True
    subtitle_run.font.size = Pt(12)
    subtitle_run.font.color.rgb = RGBColor(72, 101, 129)

    metadata = document.add_table(rows=3, cols=3)
    metadata.style = "Table Grid"
    metadata.alignment = WD_TABLE_ALIGNMENT.CENTER
    values = (
        ("Committee Name", committee.name),
        ("Authority Level", _format_label(committee.authority_level)),
        ("Committee Type", _format_label(committee.committee_type)),
        ("Meeting Title", meeting_title or "Meeting Agenda"),
        ("Meeting Date", meeting_date),
        ("Generated At UTC", _format_datetime(generated_at)),
        ("Generated By User ID", generated_by_user_id),
        ("Queue Scope", _queue_scope_text(queue_scope)),
        ("Report Type", REPORT_TITLE),
    )
    for cell, (label, value) in zip(
        (cell for row in metadata.rows for cell in row.cells),
        values,
        strict=True,
    ):
        cell.text = ""
        paragraph = cell.paragraphs[0]
        paragraph.add_run(f"{label}: ").bold = True
        paragraph.add_run(_safe_text(value))

    footer = document.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run(f"{REPORT_TITLE} | {committee.name}")


def generate_committee_meeting_pack_docx(
    db: Session,
    *,
    committee_id: uuid.UUID,
    generated_by_user_id: uuid.UUID,
    output_dir: Path | str,
    meeting_title: str | None = None,
    meeting_date: date | None = None,
) -> Path:
    committee, _user = _validate_generator(
        db,
        committee_id=committee_id,
        generated_by_user_id=generated_by_user_id,
    )
    try:
        risks = get_decision_queue_for_committee(
            db,
            committee_id=committee.id,
        )
    except DecisionQueueBusinessRuleError as exc:
        raise CommitteeMeetingPackBusinessRuleError(str(exc)) from exc

    risk_ids = [risk.id for risk in risks]
    assessments = (
        list(
            db.scalars(
                select(RiskAssessment).where(
                    RiskAssessment.risk_record_id.in_(risk_ids)
                )
            ).all()
        )
        if risk_ids
        else []
    )
    actions = (
        list(
            db.scalars(
                select(RiskAction).where(RiskAction.risk_record_id.in_(risk_ids))
            ).all()
        )
        if risk_ids
        else []
    )
    evidence_items = (
        list(
            db.scalars(
                select(RiskEvidence).where(
                    RiskEvidence.risk_record_id.in_(risk_ids),
                    RiskEvidence.is_active.is_(True),
                )
            ).all()
        )
        if risk_ids
        else []
    )
    monitoring_reviews = (
        list(
            db.scalars(
                select(RiskMonitoringReview).where(
                    RiskMonitoringReview.risk_record_id.in_(risk_ids),
                    RiskMonitoringReview.is_active.is_(True),
                    RiskMonitoringReview.status.in_(OPEN_MONITORING_STATUSES),
                )
            ).all()
        )
        if risk_ids
        else []
    )

    latest_assessments = _latest_assessments(assessments)
    actions_by_risk = _group_by_risk(actions)
    evidence_by_risk = _group_by_risk(evidence_items)
    monitoring_by_risk = _group_by_risk(monitoring_reviews)
    board_ids = {risk.board_of_origin_id for risk in risks if risk.board_of_origin_id}
    boards = {
        board.id: board
        for board in db.scalars(
            select(Committee).where(Committee.id.in_(board_ids))
        ).all()
    }
    queue_scope = get_decision_queue_scope(committee)
    generated_at = datetime.now(timezone.utc)

    output_path = Path(output_dir)
    output_path.mkdir(parents=True, exist_ok=True)
    dated_suffix = (meeting_date or generated_at.date()).isoformat()
    file_path = output_path / (
        f"{_safe_filename(committee.name)}_{dated_suffix}_meeting_pack.docx"
    )

    document = Document()
    _set_document_defaults(document)
    _add_header(
        document,
        committee=committee,
        generated_by_user_id=generated_by_user_id,
        generated_at=generated_at,
        queue_scope=queue_scope,
        meeting_title=meeting_title,
        meeting_date=meeting_date,
    )

    _add_section_heading(document, "Section 1 - Meeting Agenda")
    for item in MEETING_AGENDA_ITEMS:
        document.add_paragraph(item, style="List Number")

    open_actions_by_risk = {
        risk.id: [
            action
            for action in actions_by_risk.get(risk.id, [])
            if is_action_open_for_alerts(action)
        ]
        for risk in risks
    }
    overdue_risk_count = sum(
        any(get_risk_action_due_status(action) == "OVERDUE" for action in risk_actions)
        for risk_actions in open_actions_by_risk.values()
    )
    active_monitoring_risk_count = sum(
        bool(monitoring_by_risk.get(risk.id)) for risk in risks
    )
    evidence_risk_count = sum(bool(evidence_by_risk.get(risk.id)) for risk in risks)

    _add_section_heading(document, "Section 2 - Decision Queue Summary")
    _add_key_value_paragraph(
        document, "Total risks awaiting committee decision", len(risks)
    )
    _add_key_value_paragraph(
        document, "Number with overdue actions", overdue_risk_count
    )
    _add_key_value_paragraph(
        document, "Number with active monitoring reviews", active_monitoring_risk_count
    )
    _add_key_value_paragraph(
        document, "Number with evidence attached", evidence_risk_count
    )
    workflow_counts = Counter(_format_label(risk.workflow_status) for risk in risks)
    domain_counts = Counter(_format_label(risk.domain) for risk in risks)
    _add_table(
        document,
        ["Workflow Status", "Risk Count"],
        [[status, count] for status, count in sorted(workflow_counts.items())],
    )
    _add_table(
        document,
        ["Risk Domain", "Risk Count"],
        [[domain, count] for domain, count in sorted(domain_counts.items())],
    )

    _add_section_heading(document, "Section 3 - Risks Awaiting Committee Review")
    if risks:
        risk_rows = []
        for risk in risks:
            initial = latest_assessments.get((risk.id, RiskAssessmentType.INITIAL))
            residual = latest_assessments.get((risk.id, RiskAssessmentType.RESIDUAL))
            open_actions = open_actions_by_risk[risk.id]
            risk_rows.append(
                [
                    _risk_display_id(risk),
                    _format_label(risk.domain),
                    boards[risk.board_of_origin_id].name
                    if risk.board_of_origin_id in boards
                    else "Not recorded.",
                    _format_label(risk.workflow_status),
                    _format_label(risk.lifecycle_status),
                    risk.problem_description,
                    (residual or initial).risk_level if residual or initial else None,
                    residual.risk_level if residual else None,
                    len(open_actions),
                    sum(
                        get_risk_action_due_status(action) == "OVERDUE"
                        for action in open_actions
                    ),
                    len(evidence_by_risk.get(risk.id, [])),
                    _monitoring_summary(monitoring_by_risk.get(risk.id, [])),
                    _format_datetime(risk.updated_at),
                ]
            )
        _add_table(
            document,
            [
                "Risk ID",
                "Domain",
                "Board of Origin",
                "Workflow Status",
                "Lifecycle Status",
                "Problem Description",
                "Latest Risk Level",
                "Residual Risk Level",
                "Open Actions",
                "Overdue Actions",
                "Evidence Count",
                "Monitoring Status",
                "Updated At",
            ],
            risk_rows,
        )
    else:
        document.add_paragraph("No risks are currently in this committee Decision Queue.")

    _add_section_heading(document, "Section 4 - Individual Risk Briefs")
    if not risks:
        document.add_paragraph("No individual risk briefs are required for this agenda.")
    for risk in risks:
        document.add_heading(_risk_display_id(risk), level=2)
        initial = latest_assessments.get((risk.id, RiskAssessmentType.INITIAL))
        residual = latest_assessments.get((risk.id, RiskAssessmentType.RESIDUAL))
        open_actions = open_actions_by_risk[risk.id]
        overdue_actions = [
            action
            for action in open_actions
            if get_risk_action_due_status(action) == "OVERDUE"
        ]
        action_titles = ", ".join(action.title for action in open_actions[:5])
        _add_key_value_paragraph(document, "Risk ID", _risk_display_id(risk))
        _add_key_value_paragraph(
            document, "Problem Description", risk.problem_description
        )
        _add_key_value_paragraph(document, "System Scope", risk.system_scope)
        _add_key_value_paragraph(document, "Central Event", risk.central_event)
        _add_key_value_paragraph(document, "Hazard Statement", risk.hazard_statement)
        _add_key_value_paragraph(
            document, "Workflow Status", _format_label(risk.workflow_status)
        )
        _add_key_value_paragraph(
            document, "Lifecycle Status", _format_label(risk.lifecycle_status)
        )
        _add_key_value_paragraph(
            document,
            "Board of Origin / Originating Committee",
            boards[risk.board_of_origin_id].name
            if risk.board_of_origin_id in boards
            else None,
        )
        _add_key_value_paragraph(
            document, "Current Assessment Summary", _assessment_summary(initial)
        )
        _add_key_value_paragraph(
            document, "Residual Assessment Summary", _assessment_summary(residual)
        )
        _add_key_value_paragraph(
            document,
            "Open Actions Summary",
            f"{len(open_actions)} open; {len(overdue_actions)} overdue"
            + (f"; {action_titles}" if action_titles else ""),
        )
        _add_key_value_paragraph(
            document,
            "Evidence Metadata Count",
            len(evidence_by_risk.get(risk.id, [])),
        )
        _add_key_value_paragraph(
            document,
            "Monitoring Review Status",
            _monitoring_summary(monitoring_by_risk.get(risk.id, [])),
        )
        _add_key_value_paragraph(
            document,
            "Suggested Committee Focus",
            "Decision required based on current workflow status.",
        )

    _add_section_heading(document, "Section 5 - Action Follow-up")
    open_action_rows = [
        [
            _risk_display_id(risk),
            action.title,
            action.action_owner_user_id,
            _format_datetime(action.due_date),
            _format_label(get_risk_action_due_status(action)),
            _format_label(action.status),
        ]
        for risk in risks
        for action in open_actions_by_risk[risk.id]
    ]
    if open_action_rows:
        _add_table(
            document,
            [
                "Risk ID",
                "Action Title",
                "Action Owner",
                "Due Date",
                "Due Status",
                "Action Status",
            ],
            open_action_rows,
        )
    else:
        document.add_paragraph("No open Risk Actions require follow-up.")

    _add_section_heading(document, "Section 6 - Monitoring Follow-up")
    monitoring_rows = [
        [
            _risk_display_id(risk),
            _format_label(review.status),
            _format_datetime(review.next_review_date),
            review.review_frequency,
            _format_label(review.review_outcome) if review.review_outcome else None,
            _format_datetime(review.last_reviewed_at),
        ]
        for risk in risks
        for review in monitoring_by_risk.get(risk.id, [])
    ]
    if monitoring_rows:
        _add_table(
            document,
            [
                "Risk ID",
                "Monitoring Status",
                "Next Review Date",
                "Review Frequency",
                "Review Outcome",
                "Last Reviewed At",
            ],
            monitoring_rows,
        )
    else:
        document.add_paragraph("No active monitoring reviews require follow-up.")

    _add_section_heading(document, "Section 7 - Evidence Summary")
    evidence_rows = [
        [
            _risk_display_id(risk),
            evidence.original_filename,
            evidence.description,
            _format_datetime(evidence.uploaded_at),
            "Active",
        ]
        for risk in risks
        for evidence in evidence_by_risk.get(risk.id, [])
    ]
    if evidence_rows:
        _add_table(
            document,
            ["Risk ID", "Evidence File Name", "Description", "Uploaded At", "Status"],
            evidence_rows,
        )
    else:
        document.add_paragraph("No active evidence metadata is attached to queued risks.")
    document.add_paragraph("Evidence files are not embedded in this meeting pack.")

    _add_section_heading(document, "Section 8 - Meeting Notes Placeholder")
    document.add_paragraph(
        "Meeting notes and decisions may be recorded in the system after committee review."
    )
    for _ in range(4):
        document.add_paragraph("____________________________________________________________________")

    _add_section_heading(document, "Section 9 - Disclaimer")
    document.add_paragraph(
        "This meeting pack is generated from system records. It supports committee "
        "review and does not replace formal committee decision entry, accountable "
        "manager acceptance, or required regulatory documentation."
    )

    try:
        document.save(file_path)
    except Exception as exc:
        raise CommitteeMeetingPackBusinessRuleError(
            "Failed to generate Committee Meeting Pack DOCX"
        ) from exc
    return file_path

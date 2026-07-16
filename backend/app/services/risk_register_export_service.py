import csv
import json
import uuid
from dataclasses import asdict, dataclass
from datetime import date, datetime, timezone
from enum import Enum
from pathlib import Path
from typing import Any

from docx import Document
from docx.document import Document as DocumentObject
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from sqlalchemy import func, select
from sqlalchemy.orm import Session

from app.models.committee import Committee
from app.models.enums import (
    RiskActionStatus,
    RiskMonitoringStatus,
)
from app.models.risk import (
    RiskAction,
    RiskAssessment,
    RiskEvidence,
    RiskMonitoringReview,
    RiskRecord,
)
from app.schemas.risk_search import RiskRecordListFilters
from app.services.risk_service import (
    RiskRecordBusinessRuleError,
    list_authorized_risk_records,
)

DEFAULT_RISK_REGISTER_EXPORT_OUTPUT_DIR = Path("generated_reports")
RISK_REGISTER_EXPORT_COLUMNS = [
    "risk_record_id",
    "risk_id",
    "problem_description",
    "domain",
    "source_trigger",
    "board_of_origin_id",
    "board_of_origin_name",
    "board_of_origin_authority_level",
    "workflow_status",
    "lifecycle_status",
    "owner_user_id",
    "created_by_user_id",
    "latest_assessment_type",
    "latest_risk_level",
    "latest_assessment_date",
    "open_action_count",
    "overdue_action_count",
    "monitoring_status",
    "evidence_count",
    "created_at",
    "updated_at",
    "is_active",
]
NAVY = "17365D"
LIGHT_BLUE = "DCE6F1"


class RiskRegisterExportBusinessRuleError(ValueError):
    pass


@dataclass(frozen=True)
class RiskRegisterExportRow:
    risk_record_id: str
    risk_id: str
    problem_description: str
    domain: str
    source_trigger: str
    board_of_origin_id: str
    board_of_origin_name: str
    board_of_origin_authority_level: str
    workflow_status: str
    lifecycle_status: str
    owner_user_id: str
    created_by_user_id: str
    latest_assessment_type: str
    latest_risk_level: str
    latest_assessment_date: str
    open_action_count: int
    overdue_action_count: int
    monitoring_status: str
    evidence_count: int
    created_at: str
    updated_at: str
    is_active: bool


def _utc_now() -> datetime:
    return datetime.now(timezone.utc)


def _timestamp_for_filename(value: datetime) -> str:
    return value.strftime("%Y%m%d_%H%M%S")


def _ensure_output_dir(output_dir: Path | str) -> Path:
    output_path = Path(output_dir)
    output_path.mkdir(parents=True, exist_ok=True)
    return output_path


def _json_default(value: Any) -> str:
    if isinstance(value, datetime):
        return value.isoformat()
    if isinstance(value, date):
        return value.isoformat()
    if isinstance(value, uuid.UUID):
        return str(value)
    if isinstance(value, Enum):
        return value.value
    return str(value)


def _safe_text(value: Any, *, blank: str = "") -> str:
    if value is None:
        return blank
    if isinstance(value, Enum):
        return value.value
    if isinstance(value, uuid.UUID):
        return str(value)
    if isinstance(value, datetime):
        utc_value = value.astimezone(timezone.utc) if value.tzinfo else value
        return f"{utc_value.strftime('%Y-%m-%d %H:%M:%S')} UTC"
    if isinstance(value, date):
        return value.isoformat()
    if isinstance(value, dict | list):
        return json.dumps(value, default=_json_default, ensure_ascii=False)
    return str(value)


def _format_filter(value: Any) -> str:
    text = _safe_text(value)
    return text if text else "All"


def _truncate(value: str, max_length: int) -> str:
    if len(value) <= max_length:
        return value
    return f"{value[: max_length - 3]}..."


def get_authorized_risk_register_export_rows(
    db: Session,
    *,
    requested_by_user_id: uuid.UUID | None,
    filters: RiskRecordListFilters,
) -> list[RiskRegisterExportRow]:
    try:
        risk_records = list_authorized_risk_records(
            db,
            requested_by_user_id=requested_by_user_id,
            include_archived=filters.include_archived,
            filters=filters,
        )
    except RiskRecordBusinessRuleError as exc:
        raise RiskRegisterExportBusinessRuleError(str(exc)) from exc

    if not risk_records:
        return []

    risk_ids = [risk.id for risk in risk_records]
    committees = _get_committees_by_id(db, risk_records)
    latest_assessments = _get_latest_assessments_by_risk_id(db, risk_ids)
    open_action_counts = _get_open_action_counts(db, risk_ids)
    overdue_action_counts = _get_overdue_action_counts(db, risk_ids)
    monitoring_statuses = _get_monitoring_statuses(db, risk_ids)
    evidence_counts = _get_evidence_counts(db, risk_ids)

    rows: list[RiskRegisterExportRow] = []
    for risk in risk_records:
        committee = (
            committees.get(risk.board_of_origin_id)
            if risk.board_of_origin_id is not None
            else None
        )
        latest_assessment = latest_assessments.get(risk.id)
        rows.append(
            RiskRegisterExportRow(
                risk_record_id=str(risk.id),
                risk_id=risk.risk_id or "",
                problem_description=risk.problem_description,
                domain=_safe_text(risk.domain),
                source_trigger=risk.source_trigger or "",
                board_of_origin_id=_safe_text(risk.board_of_origin_id),
                board_of_origin_name=committee.name if committee is not None else "",
                board_of_origin_authority_level=(
                    _safe_text(committee.authority_level)
                    if committee is not None
                    else ""
                ),
                workflow_status=_safe_text(risk.workflow_status),
                lifecycle_status=_safe_text(risk.lifecycle_status),
                owner_user_id=_safe_text(risk.owner_user_id),
                created_by_user_id=_safe_text(risk.created_by_user_id),
                latest_assessment_type=(
                    _safe_text(latest_assessment.assessment_type)
                    if latest_assessment is not None
                    else ""
                ),
                latest_risk_level=(
                    latest_assessment.risk_level
                    if latest_assessment is not None
                    else "Not assessed"
                ),
                latest_assessment_date=(
                    _safe_text(latest_assessment.assessed_at)
                    if latest_assessment is not None
                    else ""
                ),
                open_action_count=open_action_counts.get(risk.id, 0),
                overdue_action_count=overdue_action_counts.get(risk.id, 0),
                monitoring_status=monitoring_statuses.get(risk.id, "Not monitored"),
                evidence_count=evidence_counts.get(risk.id, 0),
                created_at=_safe_text(risk.created_at),
                updated_at=_safe_text(risk.updated_at),
                is_active=risk.is_active,
            )
        )

    return rows


def export_risk_register_csv(
    db: Session,
    *,
    requested_by_user_id: uuid.UUID | None,
    filters: RiskRecordListFilters,
    output_dir: Path | str = DEFAULT_RISK_REGISTER_EXPORT_OUTPUT_DIR,
) -> Path:
    rows = get_authorized_risk_register_export_rows(
        db,
        requested_by_user_id=requested_by_user_id,
        filters=filters,
    )
    generated_at = _utc_now()
    file_path = (
        _ensure_output_dir(output_dir)
        / f"risk_register_export_{_timestamp_for_filename(generated_at)}.csv"
    )

    with file_path.open("w", encoding="utf-8", newline="") as csv_file:
        writer = csv.DictWriter(csv_file, fieldnames=RISK_REGISTER_EXPORT_COLUMNS)
        writer.writeheader()
        for row in rows:
            writer.writerow(asdict(row))

    return file_path


def export_risk_register_docx(
    db: Session,
    *,
    requested_by_user_id: uuid.UUID | None,
    filters: RiskRecordListFilters,
    output_dir: Path | str = DEFAULT_RISK_REGISTER_EXPORT_OUTPUT_DIR,
) -> Path:
    rows = get_authorized_risk_register_export_rows(
        db,
        requested_by_user_id=requested_by_user_id,
        filters=filters,
    )
    generated_at = _utc_now()
    file_path = (
        _ensure_output_dir(output_dir)
        / f"risk_register_export_{_timestamp_for_filename(generated_at)}.docx"
    )

    document = Document()
    _set_document_defaults(document)
    _add_header(
        document,
        generated_at=generated_at,
        generated_by_user_id=requested_by_user_id,
        filters=filters,
        record_count=len(rows),
    )

    document.add_heading("Export Scope", level=1)
    document.add_paragraph(
        "This Risk Register Export contains risk records authorized for the "
        "requesting user and filtered according to the export criteria."
    )

    document.add_heading("Register Summary", level=1)
    _add_summary_table(document, rows)

    document.add_heading("Risk Register", level=1)
    _add_register_table(document, rows)

    document.add_heading("Risk Details", level=1)
    _add_detail_sections(document, rows)

    document.add_heading("Disclaimer", level=1)
    document.add_paragraph(
        "This Risk Register Export is generated from system records. It supports "
        "SMS governance, audit preparation, committee preparation, management "
        "review, and offline analysis. It does not replace formal risk assessment, "
        "committee decision entry, accountable manager acceptance, or required "
        "regulatory documentation."
    )

    document.save(file_path)
    return file_path


def _get_committees_by_id(
    db: Session, risk_records: list[RiskRecord]
) -> dict[uuid.UUID, Committee]:
    committee_ids = {
        risk.board_of_origin_id
        for risk in risk_records
        if risk.board_of_origin_id is not None
    }
    if not committee_ids:
        return {}
    committees = db.scalars(select(Committee).where(Committee.id.in_(committee_ids)))
    return {committee.id: committee for committee in committees}


def _get_latest_assessments_by_risk_id(
    db: Session, risk_ids: list[uuid.UUID]
) -> dict[uuid.UUID, RiskAssessment]:
    assessments = db.scalars(
        select(RiskAssessment)
        .where(RiskAssessment.risk_record_id.in_(risk_ids))
        .order_by(
            RiskAssessment.risk_record_id.asc(),
            RiskAssessment.assessed_at.desc(),
            RiskAssessment.created_at.desc(),
        )
    ).all()

    latest: dict[uuid.UUID, RiskAssessment] = {}
    for assessment in assessments:
        latest.setdefault(assessment.risk_record_id, assessment)
    return latest


def _get_open_action_counts(
    db: Session, risk_ids: list[uuid.UUID]
) -> dict[uuid.UUID, int]:
    return dict(
        db.execute(
            select(RiskAction.risk_record_id, func.count(RiskAction.id))
            .where(
                RiskAction.risk_record_id.in_(risk_ids),
                RiskAction.status.in_(
                    [RiskActionStatus.OPEN, RiskActionStatus.IN_PROGRESS]
                ),
            )
            .group_by(RiskAction.risk_record_id)
        ).all()
    )


def _get_overdue_action_counts(
    db: Session, risk_ids: list[uuid.UUID]
) -> dict[uuid.UUID, int]:
    return dict(
        db.execute(
            select(RiskAction.risk_record_id, func.count(RiskAction.id))
            .where(
                RiskAction.risk_record_id.in_(risk_ids),
                RiskAction.status.in_(
                    [RiskActionStatus.OPEN, RiskActionStatus.IN_PROGRESS]
                ),
                RiskAction.due_date < date.today(),
            )
            .group_by(RiskAction.risk_record_id)
        ).all()
    )


def _get_monitoring_statuses(
    db: Session, risk_ids: list[uuid.UUID]
) -> dict[uuid.UUID, str]:
    reviews = db.execute(
        select(RiskMonitoringReview.risk_record_id, RiskMonitoringReview.status)
        .where(
            RiskMonitoringReview.risk_record_id.in_(risk_ids),
            RiskMonitoringReview.is_active.is_(True),
        )
    ).all()

    status_sets: dict[uuid.UUID, set[RiskMonitoringStatus]] = {}
    for risk_record_id, status in reviews:
        status_sets.setdefault(risk_record_id, set()).add(status)

    priority = [
        RiskMonitoringStatus.OVERDUE,
        RiskMonitoringStatus.DUE,
        RiskMonitoringStatus.ACTIVE,
        RiskMonitoringStatus.CLOSED,
    ]
    summaries: dict[uuid.UUID, str] = {}
    for risk_record_id, statuses in status_sets.items():
        summaries[risk_record_id] = next(
            (status.value for status in priority if status in statuses),
            "Not monitored",
        )
    return summaries


def _get_evidence_counts(
    db: Session, risk_ids: list[uuid.UUID]
) -> dict[uuid.UUID, int]:
    return dict(
        db.execute(
            select(RiskEvidence.risk_record_id, func.count(RiskEvidence.id))
            .where(
                RiskEvidence.risk_record_id.in_(risk_ids),
                RiskEvidence.is_active.is_(True),
            )
            .group_by(RiskEvidence.risk_record_id)
        ).all()
    )


def _set_document_defaults(document: DocumentObject) -> None:
    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)

    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(8.5)
    styles["Title"].font.name = "Arial"
    styles["Title"].font.size = Pt(24)
    styles["Title"].font.color.rgb = RGBColor(23, 54, 93)
    for style_name, font_size in (("Heading 1", 14), ("Heading 2", 11)):
        style = styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(font_size)
        style.font.color.rgb = RGBColor(23, 54, 93)


def _add_header(
    document: DocumentObject,
    *,
    generated_at: datetime,
    generated_by_user_id: uuid.UUID | None,
    filters: RiskRecordListFilters,
    record_count: int,
) -> None:
    title = document.add_heading("Risk Register Export", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run("SMS Risk Management Process Tool")
    subtitle_run.bold = True
    subtitle_run.font.size = Pt(12)
    subtitle_run.font.color.rgb = RGBColor(72, 101, 129)

    metadata_values = [
        ("Export Type", "Controlled Export"),
        ("Generated At UTC", _safe_text(generated_at)),
        ("Generated By User ID", _format_filter(generated_by_user_id)),
        ("Record Count", str(record_count)),
        ("Search filter", _format_filter(filters.search)),
        ("Risk ID filter", _format_filter(filters.risk_id)),
        ("Domain filter", _format_filter(filters.domain)),
        ("Board of Origin filter", _format_filter(filters.board_of_origin_id)),
        ("Workflow Status filter", _format_filter(filters.workflow_status)),
        ("Lifecycle Status filter", _format_filter(filters.lifecycle_status)),
        ("Latest Risk Level filter", _format_filter(filters.latest_risk_level)),
        ("Include Archived", "Yes" if filters.include_archived else "No"),
        ("Sort By", filters.sort_by),
        ("Sort Direction", filters.sort_direction),
    ]
    metadata = document.add_table(rows=7, cols=2)
    metadata.style = "Table Grid"
    metadata.alignment = WD_TABLE_ALIGNMENT.CENTER

    for cell, (label, value) in zip(
        (cell for row in metadata.rows for cell in row.cells),
        metadata_values,
        strict=True,
    ):
        cell.text = ""
        paragraph = cell.paragraphs[0]
        paragraph.add_run(f"{label}: ").bold = True
        paragraph.add_run(value)

    footer = document.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.text = "Risk Register Export - Controlled Export"


def _add_summary_table(
    document: DocumentObject, rows: list[RiskRegisterExportRow]
) -> None:
    summary_values = [
        ("Total exported risks", len(rows)),
        ("Active risks", sum(1 for row in rows if row.is_active)),
        ("Archived/inactive risks", sum(1 for row in rows if not row.is_active)),
        (
            "Not assessed risks",
            sum(1 for row in rows if row.latest_risk_level == "Not assessed"),
        ),
        (
            "Risks with overdue actions",
            sum(1 for row in rows if row.overdue_action_count > 0),
        ),
        (
            "Risks under monitoring",
            sum(1 for row in rows if row.monitoring_status != "Not monitored"),
        ),
        ("Risks with evidence", sum(1 for row in rows if row.evidence_count > 0)),
    ]
    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    _style_header_cells(table.rows[0].cells, ["Metric", "Value"])
    for label, value in summary_values:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = str(value)


def _add_register_table(
    document: DocumentObject, rows: list[RiskRegisterExportRow]
) -> None:
    headers = [
        "Risk ID",
        "Domain",
        "Board of Origin",
        "Workflow Status",
        "Lifecycle Status",
        "Latest Risk Level",
        "Open Actions",
        "Overdue Actions",
        "Monitoring",
        "Evidence Count",
        "Updated",
    ]
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _style_header_cells(table.rows[0].cells, headers)

    if not rows:
        cells = table.add_row().cells
        cells[0].text = "No authorized risk records matched the export criteria."
        for cell in cells[1:]:
            cell.text = ""
        return

    for row_data in rows:
        values = [
            row_data.risk_id or row_data.risk_record_id,
            row_data.domain,
            _truncate(row_data.board_of_origin_name or "Not assigned", 55),
            row_data.workflow_status,
            row_data.lifecycle_status,
            row_data.latest_risk_level,
            str(row_data.open_action_count),
            str(row_data.overdue_action_count),
            row_data.monitoring_status,
            str(row_data.evidence_count),
            row_data.updated_at,
        ]
        row = table.add_row()
        for cell, value in zip(row.cells, values, strict=True):
            cell.text = value
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(6.5)


def _add_detail_sections(
    document: DocumentObject, rows: list[RiskRegisterExportRow]
) -> None:
    if not rows:
        document.add_paragraph("No authorized risk records matched the export criteria.")
        return

    for row in rows:
        document.add_heading(row.risk_id or row.risk_record_id, level=2)
        table = document.add_table(rows=0, cols=2)
        table.style = "Table Grid"
        detail_values = [
            ("Risk ID", row.risk_id or row.risk_record_id),
            ("Problem Description", row.problem_description),
            ("Source Trigger", row.source_trigger or "Not recorded"),
            ("Board of Origin", row.board_of_origin_name or "Not assigned"),
            ("Authority Level", row.board_of_origin_authority_level or "Not assigned"),
            ("Workflow Status", row.workflow_status),
            ("Lifecycle Status", row.lifecycle_status),
            ("Latest Risk Level", row.latest_risk_level),
            (
                "Latest Assessment Type / Date",
                (
                    f"{row.latest_assessment_type} / {row.latest_assessment_date}"
                    if row.latest_assessment_type
                    else "Not assessed"
                ),
            ),
            ("Open Actions", str(row.open_action_count)),
            ("Overdue Actions", str(row.overdue_action_count)),
            ("Monitoring Status", row.monitoring_status),
            ("Evidence Count", str(row.evidence_count)),
        ]
        for label, value in detail_values:
            cells = table.add_row().cells
            cells[0].text = label
            cells[0].paragraphs[0].runs[0].bold = True
            cells[1].text = value


def _style_header_cells(cells, headers: list[str]) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    for cell, header in zip(cells, headers, strict=True):
        cell.text = header
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        shading = cell._tc.get_or_add_tcPr()
        fill = OxmlElement("w:shd")
        fill.set(qn("w:fill"), LIGHT_BLUE)
        shading.append(fill)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(23, 54, 93)
                run.font.size = Pt(7)
